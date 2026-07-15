import base64
import io
import logging
import re
import sys
from pathlib import Path
from typing import TYPE_CHECKING, Any

from docx import Document
from docx.oxml import parse_xml
from docx.oxml import parser as docx_parser
from docx.oxml.ns import nsdecls
from lxml import etree  # type: ignore

if TYPE_CHECKING:
    from collections.abc import Iterator

    from docx.document import Document as DocumentObject
    from docx.section import Section
    from docx.table import Table, _Cell

    from app.HtmlTableLayout import TableLayout

from app.DocxMathColorPostProcess import apply_math_colors
from app.DocxReferencesPostProcess import add_table_of_contents_entries, enable_auto_update_fields

# Patch the python-docx parser to handle large XML documents (> 10MB)
# This enables the XML_PARSE_HUGE flag to avoid "Buffer size limit exceeded" errors
# when processing documents with large embedded content (e.g., base64-encoded images)
_huge_tree_parser = etree.XMLParser(remove_blank_text=True, resolve_entities=False, huge_tree=True)
_huge_tree_parser.set_element_class_lookup(docx_parser.element_class_lookup)
docx_parser.oxml_parser = _huge_tree_parser

SCHEMA = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

EMU_1_INCH = 914400  # 1 inch in docx in EMU (English Metric Units)
TWIPS_1_INCH = 1440  # 1 inch in docx in Twips (Twentieth of a Point)
DOCX_LETTER_WIDTH_EMU = 8.5 * EMU_1_INCH  # docx LETTER width = 8.5 inch
DOCX_LETTER_SIDE_MARGIN = EMU_1_INCH  # docx left & right margins = 1 inch

# Paper sizes in TWIPS (portrait orientation: width x height)
PAPER_SIZES = {
    "A5": {"width": 8419, "height": 11906},
    "A4": {"width": 11906, "height": 16838},
    "A3": {"width": 16838, "height": 23811},
    "B5": {"width": 9979, "height": 14144},
    "B4": {"width": 14144, "height": 20013},
    "JIS_B5": {"width": 10319, "height": 14572},
    "JIS_B4": {"width": 14572, "height": 20639},
    "LETTER": {"width": 12240, "height": 15840},
    "LEGAL": {"width": 12240, "height": 20160},
    "LEDGER": {"width": 15840, "height": 24480},
}
logger = logging.getLogger(__name__)

# OOXML tblPr child element order (subset of CT_TblPrBase we touch). Used to
# insert new properties at a schema-valid position so both Word and the
# stricter LibreOffice accept the output regardless of which children pandoc
# (or the inline_styles.lua table rebuild) already emitted.
_TBLPR_CHILD_ORDER = [
    "tblStyle",
    "tblpPr",
    "tblOverlap",
    "bidiVisual",
    "tblStyleRowBandSize",
    "tblStyleColBandSize",
    "tblW",
    "jc",
    "tblCellSpacing",
    "tblInd",
    "tblBorders",
    "shd",
    "tblLayout",
    "tblCellMar",
    "tblLook",
    "tblCaption",
    "tblDescription",
]


def process(docx_bytes: bytes, paper_size: str | None = None, orientation: str | None = None, table_layouts: list[TableLayout] | None = None) -> bytes:
    doc = Document(io.BytesIO(docx_bytes))
    _move_header_footer_references_to_first_section(doc)
    _replace_first_paragraph_styles(doc)
    _replace_size_and_orientation(doc, paper_size, orientation)
    _replace_table_properties(doc, table_layouts)
    apply_math_colors(doc)
    _replace_image_placeholders(doc)
    _replace_link_placeholders(doc)
    add_table_of_contents_entries(doc)
    enable_auto_update_fields(doc)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


_IMG_PLACEHOLDER_RE = re.compile(r"\{\{IMG:(.*?)\}\}")
_HREF_PLACEHOLDER_RE = re.compile(r"\{\{HREF:(.*?)\}\}")

RELATIONSHIPS_SCHEMA = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"  # NOSONAR


def _replace_image_placeholders(doc: DocumentObject) -> None:
    """Replace ``{{IMG:<src>}}`` placeholders with real embedded images.

    The ``inline_styles.lua`` filter emits these markers when it rebuilds a
    styled table as raw OOXML. Images can't be embedded in raw OOXML (they
    need writer-level relationship entries), so the Lua filter writes a
    text placeholder and this function resolves it using python-docx.
    """
    body = doc.element.body
    doc_pr_id = max((int(dp.get("id", "0")) for dp in body.iter("{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}docPr")), default=0)
    for t_el in body.findall(f".//{{{SCHEMA}}}t"):
        if t_el.text is None:
            continue
        match = _IMG_PLACEHOLDER_RE.search(t_el.text)
        if not match:
            continue

        src = match.group(1)
        run_el = t_el.getparent()
        if run_el is None or not run_el.tag.endswith("}r"):
            continue

        image_bytes = _resolve_image_src(src)
        if image_bytes is None:
            # Can't resolve — leave the placeholder text as alt-text fallback
            t_el.text = t_el.text.replace(match.group(0), "[image]")
            continue

        try:
            r_id, img = doc.part.get_or_add_image(io.BytesIO(image_bytes))
            width = img.px_width * EMU_1_INCH // 96  # px to EMU at 96 dpi
            height = img.px_height * EMU_1_INCH // 96
            doc_pr_id += 1

            # Build the drawing XML
            drawing_xml = (
                f'<w:drawing {nsdecls("w", "wp", "a", "pic", "r")}>'
                f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
                f'<wp:extent cx="{width}" cy="{height}"/>'
                f'<wp:docPr id="{doc_pr_id}" name="Image"/>'
                f"<a:graphic>"
                f'<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
                f"<pic:pic>"
                f"<pic:nvPicPr>"
                f'<pic:cNvPr id="{doc_pr_id}" name="Image"/>'
                f"<pic:cNvPicPr/>"
                f"</pic:nvPicPr>"
                f"<pic:blipFill>"
                f'<a:blip r:embed="{r_id}"/>'
                f"<a:stretch><a:fillRect/></a:stretch>"
                f"</pic:blipFill>"
                f"<pic:spPr>"
                f"<a:xfrm>"
                f'<a:off x="0" y="0"/>'
                f'<a:ext cx="{width}" cy="{height}"/>'
                f"</a:xfrm>"
                f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
                f"</pic:spPr>"
                f"</pic:pic>"
                f"</a:graphicData>"
                f"</a:graphic>"
                f"</wp:inline>"
                f"</w:drawing>"
            )
            drawing_el = parse_xml(drawing_xml)

            # Replace the text run with the drawing
            run_el.remove(t_el)
            run_el.append(drawing_el)

            logger.debug(f"Replaced image placeholder with embedded image ({img.px_width}x{img.px_height})")
        except Exception as e:
            logger.warning(f"Could not embed image from placeholder: {e}")
            t_el.text = t_el.text.replace(match.group(0), "[image]")


def _resolve_image_src(src: str) -> bytes | None:
    """Resolve an image src to bytes. Supports data: URIs."""
    if src.startswith("data:"):
        # data:image/gif;base64,AAAA...
        match = re.match(r"data:[^;]+;base64,(.*)", src)
        if match:
            try:
                return base64.b64decode(match.group(1))
            except Exception:
                logger.warning("Failed to decode base64 image data")
                return None
    if src:
        logger.warning("Unsupported image src scheme (only data: URIs are supported): %s", src[:80])
    return None


def _replace_link_placeholders(doc: DocumentObject) -> None:
    """Replace ``{{HREF:<url>}}`` placeholders in hyperlink tooltips with real relationships.

    The ``inline_styles.lua`` filter emits ``<w:hyperlink w:tooltip="{{HREF:url}}">``
    when it rebuilds styled tables as raw OOXML. This function registers the
    URL as a hyperlink relationship and sets the correct ``r:id``.
    """
    ns_w = f"{{{SCHEMA}}}"
    ns_r = f"{{{RELATIONSHIPS_SCHEMA}}}"
    body = doc.element.body

    for hyperlink in body.findall(f".//{ns_w}hyperlink"):
        tooltip = hyperlink.get(f"{ns_w}tooltip", "")
        match = _HREF_PLACEHOLDER_RE.search(tooltip)
        if not match:
            continue

        url = match.group(1)
        try:
            r_id = doc.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)  # NOSONAR
            hyperlink.set(f"{ns_r}id", r_id)
            hyperlink.attrib.pop(f"{ns_w}tooltip", None)
            logger.debug(f"Resolved hyperlink placeholder to {url} (r:id={r_id})")
        except Exception as e:
            logger.warning(f"Could not resolve hyperlink placeholder: {e}")
            hyperlink.attrib.pop(f"{ns_w}tooltip", None)


def _replace_first_paragraph_styles(doc: DocumentObject) -> None:
    """Replace pandoc's "First Paragraph" style with "Body Text".

    Pandoc's DOCX writer automatically assigns "First Paragraph" to the first
    paragraph after every heading.  This creates visual inconsistency because
    only that paragraph differs in style from subsequent ones.  Normalizing all
    such paragraphs to "Body Text" gives a uniform look.
    """
    for paragraph in doc.element.body.iter(f"{{{SCHEMA}}}p"):
        p_pr = paragraph.find(f"{{{SCHEMA}}}pPr")
        if p_pr is None:
            continue
        p_style = p_pr.find(f"{{{SCHEMA}}}pStyle")
        if p_style is not None and p_style.get(f"{{{SCHEMA}}}val") == "FirstParagraph":
            p_style.set(f"{{{SCHEMA}}}val", "BodyText")


def _replace_size_and_orientation(doc: DocumentObject, paper_size: str | None = None, orientation: str | None = None) -> None:
    # If both parameters are None, no modifications needed
    if paper_size is None and orientation is None:
        return

    for section in doc.sections:
        sect_pr = section._sectPr
        pg_sz = sect_pr.find(".//w:pgSz", namespaces={"w": SCHEMA})

        if paper_size is not None:
            pg_sz = _set_paper_size(sect_pr, pg_sz, paper_size, orientation)

        if orientation is not None:
            _set_orientation(sect_pr, pg_sz, orientation)


def _set_paper_size(sect_pr: Any, pg_sz: Any, paper_size: str, orientation: str | None) -> Any:
    """Set the paper size for a section."""
    # Normalize paper_size to uppercase for case-insensitive lookup
    paper_size_upper = paper_size.upper()

    # Get dimensions for the specified paper size
    if paper_size_upper not in PAPER_SIZES:
        raise ValueError(f"Unsupported paper size: {paper_size}. Supported sizes: {', '.join(PAPER_SIZES.keys())}")

    page_dims = PAPER_SIZES[paper_size_upper]
    width = page_dims["width"]
    height = page_dims["height"]

    # Get existing orientation if present (to preserve it when changing paper size)
    existing_orientation = pg_sz.get(f"{{{SCHEMA}}}orient") if pg_sz is not None else None

    # Apply existing orientation if orientation parameter is not specified
    if orientation is None and existing_orientation == "landscape":
        width, height = height, width

    # Create or update pg_sz element
    if pg_sz is None:
        pg_sz = parse_xml(f'<w:pgSz {nsdecls("w")} w:w="{width}" w:h="{height}"/>')
        sect_pr.append(pg_sz)
    else:
        pg_sz.set(f"{{{SCHEMA}}}w", str(width))
        pg_sz.set(f"{{{SCHEMA}}}h", str(height))

    # Preserve existing orientation attribute if present and orientation parameter not specified
    if orientation is None and existing_orientation is not None:
        pg_sz.set(f"{{{SCHEMA}}}orient", existing_orientation)

    return pg_sz


def _set_orientation(sect_pr: Any, pg_sz: Any, orientation: str) -> None:
    """Set the orientation for a section."""
    # Ensure pg_sz exists (use LETTER as default if missing)
    if pg_sz is None:
        page_dims = PAPER_SIZES["LETTER"]
        width = page_dims["width"]
        height = page_dims["height"]
        pg_sz = parse_xml(f'<w:pgSz {nsdecls("w")} w:w="{width}" w:h="{height}"/>')
        sect_pr.append(pg_sz)

    # Get current dimensions
    current_width = int(pg_sz.get(f"{{{SCHEMA}}}w", "0"))
    current_height = int(pg_sz.get(f"{{{SCHEMA}}}h", "0"))

    # Determine current and desired orientation
    current_is_landscape = current_width > current_height
    desired_is_landscape = orientation.lower() == "landscape"

    # Swap dimensions if orientations don't match
    if current_is_landscape != desired_is_landscape:
        pg_sz.set(f"{{{SCHEMA}}}w", str(current_height))
        pg_sz.set(f"{{{SCHEMA}}}h", str(current_width))

    # Set or remove the orient attribute
    if desired_is_landscape:
        pg_sz.set(f"{{{SCHEMA}}}orient", "landscape")
    # Remove orient attribute for portrait (it's the default)
    elif f"{{{SCHEMA}}}orient" in pg_sz.attrib:
        del pg_sz.attrib[f"{{{SCHEMA}}}orient"]


def _move_header_footer_references_to_first_section(doc: DocumentObject) -> None:
    """
    Move header/footer references from the last section to the first section.

    When Lua filters insert section breaks (e.g., for page orientation changes),
    they create new sections without header/footer references. The template's
    header/footer refs end up only in the last section.

    In OOXML, sections without explicit header/footer refs inherit from previous
    sections. By moving the refs to the first section, all subsequent sections
    will inherit them automatically. This also properly handles first/odd/even
    page header/footer configurations.

    Additionally, the <w:titlePg/> element is moved along with the references.
    This element controls the "different first page" header/footer feature in Word.
    If left in the last section while refs are in the first, Word would incorrectly
    apply "first page" headers/footers to the wrong pages.
    """
    if len(doc.sections) <= 1:
        return  # Nothing to fix if there's only one section

    first_sect_pr = doc.sections[0]._sectPr
    last_sect_pr = doc.sections[-1]._sectPr

    # Check if first section already has header/footer references
    existing_headers = first_sect_pr.findall("w:headerReference", namespaces={"w": SCHEMA})
    existing_footers = first_sect_pr.findall("w:footerReference", namespaces={"w": SCHEMA})

    if existing_headers or existing_footers:
        return  # First section already has refs, nothing to do

    # Find and move header references from last to first section
    header_refs = last_sect_pr.findall("w:headerReference", namespaces={"w": SCHEMA})
    for ref in header_refs:
        last_sect_pr.remove(ref)
        first_sect_pr.insert(0, ref)

    # Find and move footer references from last to first section
    footer_refs = last_sect_pr.findall("w:footerReference", namespaces={"w": SCHEMA})
    for ref in footer_refs:
        last_sect_pr.remove(ref)
        first_sect_pr.insert(0, ref)

    # Move titlePg element if present (controls "different first page" header/footer)
    title_pg = last_sect_pr.find("w:titlePg", namespaces={"w": SCHEMA})
    if title_pg is not None:
        last_sect_pr.remove(title_pg)
        first_sect_pr.append(title_pg)


def _replace_table_properties(doc: DocumentObject, table_layouts: list[TableLayout] | None = None) -> None:  # NOSONAR  # needed by design
    # Per-table width/alignment recovered from the HTML source (see
    # app/HtmlTableLayout.py). The list is one entry per <table> in document
    # order (depth-first, nested included) — the same order this function walks
    # tables in — so a shared iterator lines them up index-for-index. Guard on
    # an exact count match: if pandoc dropped or added a table the alignment
    # would be off, so we skip applying layouts entirely and fall back to the
    # previous 100 %/autofit default rather than mislabel tables.
    layout_iter: Iterator[TableLayout] | None = None
    if table_layouts:
        table_count = len(doc.element.body.findall(".//w:tbl", namespaces={"w": SCHEMA}))
        if table_count == len(table_layouts):
            layout_iter = iter(table_layouts)
        else:
            logger.warning("HtmlTableLayout: %d layouts for %d tables; skipping width/alignment (fallback to defaults)", len(table_layouts), table_count)

    # Group tables by their section
    for target_index, section in enumerate(doc.sections):
        max_width = _get_available_content_width_for_section(section)

        tables_in_section = []
        current_section_index = 0

        for element in doc.element.body:
            # Section break
            if element.tag.endswith("sectPr"):
                current_section_index += 1

            # Table element
            elif element.tag.endswith("tbl") and current_section_index == target_index:
                for table in doc.tables:
                    if table._element == element:
                        tables_in_section.append(table)
                        break

        # Note: This gets all tables in the document section
        for table in tables_in_section:
            _process_table(table, 0, max_width, layout_iter)


def _process_table(table: Table, parent_columns_count: int, max_width: int, layout_iter: Iterator[TableLayout] | None = None) -> None:
    tbl = table._element
    # Pull this table's layout first, before recursing into nested tables, so
    # consumption order stays depth-first and matches HtmlTableLayout.extract.
    layout = next(layout_iter, None) if layout_iter is not None else None
    columns_count = parent_columns_count + len(table.columns)
    table_properties = tbl.find(".//w:tblPr", namespaces={"w": SCHEMA})
    if table_properties is None:
        table_properties = parse_xml(f"<w:tblPr {nsdecls('w')}/>")
        tbl.insert(0, table_properties)

    _apply_table_layout(tbl, table_properties, layout, max_width)

    # Process nested tables
    for row in table.rows:
        for cell in row.cells:
            _resize_images_in_cell(cell, max_width / columns_count)
            for sub_table in cell.tables:
                _process_table(sub_table, columns_count, max_width, layout_iter)


def _clamp_twips(width_twips: int, max_width_emu: int) -> int:
    """Clamp a width in twips to the available page width (given in EMU)."""
    if max_width_emu <= 0:
        return width_twips
    max_twips = max_width_emu // 635
    if width_twips > max_twips:
        logger.debug(f"Clamped table width from {width_twips} to {max_twips} twips")
        return max_twips
    return width_twips


def _clamp_existing_fixed_width(tbl: Any, table_properties: Any, max_width: int) -> None:
    """Clamp a Lua-filter-set dxa width to the page width if it overflows."""
    if max_width <= 0:
        return
    tbl_w = table_properties.find("w:tblW", namespaces={"w": SCHEMA})
    current = int(tbl_w.get(f"{{{SCHEMA}}}w", "0"))
    clamped = _clamp_twips(current, max_width)
    if clamped < current:
        tbl_w.set(f"{{{SCHEMA}}}w", str(clamped))
        _rescale_table_grid(tbl, clamped)


def _resolve_layout_width(layout: TableLayout | None) -> tuple[str, int, bool]:
    """Extract width parameters from an HtmlTableLayout, or return defaults."""
    if layout is not None and layout.width_type is not None and layout.width_value is not None:
        return layout.width_type, layout.width_value, layout.width_type == "dxa"
    return "pct", 5000, False


def _apply_table_layout(tbl: Any, table_properties: Any, layout: TableLayout | None, max_width: int = 0) -> None:
    """Write width, alignment and indent onto a table's <w:tblPr>."""
    has_layout = layout is not None and layout.width_type is not None and layout.width_value is not None

    if not has_layout and _has_existing_fixed_width(table_properties):
        _clamp_existing_fixed_width(tbl, table_properties, max_width)
        return

    width_type, width_value, use_fixed_layout = _resolve_layout_width(layout)

    if use_fixed_layout:
        width_value = _clamp_twips(width_value, max_width)
        _rescale_table_grid(tbl, width_value)

    _set_tblpr_child(table_properties, parse_xml(f'<w:tblW {nsdecls("w")} w:w="{width_value}" w:type="{width_type}"/>'))

    layout_type = "fixed" if use_fixed_layout else "autofit"
    _set_tblpr_child(table_properties, parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="{layout_type}"/>'))

    if layout is not None and layout.jc is not None:
        _set_tblpr_child(table_properties, parse_xml(f'<w:jc {nsdecls("w")} w:val="{layout.jc}"/>'))

    if layout is not None and layout.indent_twips is not None:
        _set_tblpr_child(table_properties, parse_xml(f'<w:tblInd {nsdecls("w")} w:w="{layout.indent_twips}" w:type="dxa"/>'))


def _has_existing_fixed_width(table_properties: Any) -> bool:
    """Return True if the table already has a fixed (dxa) width from the Lua filter."""
    tbl_w = table_properties.find("w:tblW", namespaces={"w": SCHEMA})
    if tbl_w is not None and tbl_w.get(f"{{{SCHEMA}}}type") == "dxa":
        w = tbl_w.get(f"{{{SCHEMA}}}w", "0")
        return int(w) > 0
    return False


def _set_tblpr_child(table_properties: Any, new_child: Any) -> None:
    """Replace any existing child with the same tag and insert at a schema-valid position."""
    local_name = etree.QName(new_child).localname
    for existing in table_properties.findall(f"w:{local_name}", namespaces={"w": SCHEMA}):
        table_properties.remove(existing)

    order = _TBLPR_CHILD_ORDER.index(local_name)
    for child in table_properties:
        child_local = etree.QName(child).localname
        if child_local in _TBLPR_CHILD_ORDER and _TBLPR_CHILD_ORDER.index(child_local) > order:
            child.addprevious(new_child)
            return
    table_properties.append(new_child)


def _rescale_table_grid(tbl: Any, target_twips: int) -> None:
    """Scale the <w:tblGrid> column widths so they sum to target_twips.

    Under fixed layout Word derives the rendered table width from the grid
    column widths (not <w:tblW>), so an absolute table width only takes effect
    once the grid is rescaled to match it. Proportional scaling preserves the
    relative column sizing pandoc emitted; a zero/absent grid is distributed
    evenly.
    """
    grid = tbl.find("w:tblGrid", namespaces={"w": SCHEMA})
    if grid is None:
        return
    columns = grid.findall("w:gridCol", namespaces={"w": SCHEMA})
    if not columns:
        return

    width_attr = f"{{{SCHEMA}}}w"
    widths = [int(col.get(width_attr) or 0) for col in columns]
    total = sum(widths)

    if total <= 0:
        even = max(1, target_twips // len(columns))
        for col in columns:
            col.set(width_attr, str(even))
        return

    for col, width in zip(columns, widths, strict=False):
        col.set(width_attr, str(max(1, round(width * target_twips / total))))


def _get_available_content_width_for_section(section: Section) -> int:
    # Provide alternative 'Letter' paper size params in case if they were not set explicitly in the document
    page_width = section.page_width or DOCX_LETTER_WIDTH_EMU
    left_margin = section.left_margin or DOCX_LETTER_SIDE_MARGIN
    right_margin = section.right_margin or DOCX_LETTER_SIDE_MARGIN
    return int(page_width - left_margin - right_margin)


def _resize_images_in_cell(cell: _Cell, max_image_width: float) -> None:
    cell_xml = cell._tc.xml
    # ruff: noqa: S320
    # Use huge_tree parser to handle cells with large content (e.g., base64-encoded images > 10MB)
    tree = etree.fromstring(cell_xml, docx_parser.oxml_parser)

    # Find all <wp:extent> elements that define image size
    extent_elements = tree.findall(
        ".//wp:extent",
        {"wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"},
    )

    modified = False  # Track if any image was resized

    for extent in extent_elements:
        width = int(extent.attrib["cx"])
        height = int(extent.attrib["cy"])
        logger.debug(f"Image found, size: {width} x {height}")

        # Resize only if width exceeds max_image_width
        if width > max_image_width:
            scale_factor = max_image_width / width
            new_width = int(max_image_width)
            new_height = int(height * scale_factor)  # Maintain aspect ratio

            # Apply new size
            extent.set("cx", str(new_width))
            extent.set("cy", str(new_height))

            logger.debug(f"Resized to: {new_width} x {new_height}")
            modified = True

    # If any modification was made, update the cell XML
    if modified:
        cell._tc.clear_content()
        for child in tree.iterchildren():
            cell._tc.append(child)


# Just for manual test purposes. Accepts path to docx to process.
def main() -> int:
    MIN_ARGS = 2  # script name + docx path
    MAX_ARGS = 4  # script name + docx path + paper_size + orientation
    DOCX_PATH_ARG_INDEX = 1
    PAPER_SIZE_ARG_INDEX = 2
    ORIENTATION_ARG_INDEX = 3

    if not (MIN_ARGS <= len(sys.argv) <= MAX_ARGS):
        logger.info("Usage: <path_to_docx> [paper_size] [orientation]")
        return 1

    docx_path = Path(sys.argv[DOCX_PATH_ARG_INDEX]).resolve()
    base_dir = Path.cwd().resolve()
    if not docx_path.is_relative_to(base_dir):
        logger.error(f"Refusing to access path outside the working directory: {docx_path}")
        return 1

    paper_size = sys.argv[PAPER_SIZE_ARG_INDEX] if len(sys.argv) > PAPER_SIZE_ARG_INDEX and sys.argv[PAPER_SIZE_ARG_INDEX] != "None" else None
    orientation = sys.argv[ORIENTATION_ARG_INDEX] if len(sys.argv) > ORIENTATION_ARG_INDEX and sys.argv[ORIENTATION_ARG_INDEX] != "None" else None

    with docx_path.open("rb") as docx_file_reader:
        result_bytes = process(docx_file_reader.read(), paper_size, orientation)

    with docx_path.open("wb") as docx_file_writer:
        docx_file_writer.write(result_bytes)

    logger.debug(f"Successfully modified table properties in {docx_path}")
    return 0


if __name__ == "__main__":
    main()
