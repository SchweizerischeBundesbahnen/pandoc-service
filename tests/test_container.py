import io
import logging
import time
import zipfile
from pathlib import Path
from typing import NamedTuple

import docker
import pytest
import requests
from docker.models.containers import Container
from docx import Document
from docx.shared import RGBColor
from pypdf import PdfReader

from tests.test_pptx_post_process import find_presentation_information

logger = logging.getLogger(__name__)

# Constants for Docker resources
TEST_IMAGE_NAME = "pandoc-service-test"
TEST_IMAGE_TAG = "latest"
TEST_CONTAINER_NAME = "pandoc-service-test-container"
TEST_IMAGE_FULL = f"{TEST_IMAGE_NAME}:{TEST_IMAGE_TAG}"

SOURCE_HTML = """
            <html>
                <body>
                    <h1>Simple html with an ordered list</h1>
                    <ol>
                        <li>First</li>
                        <li>Second</li>
                        <li>Third</li>
                    </ol>
                    <p>Some <b>bold German vowels ä, ö, and ü</b> at the bottom.</p>
                </body>
            </html>
            """

SOURCE_HTML_WITH_HEADINGS = """
            <html>
                <head>
                    <title>Test doc title</title>
                </head>
                <body>
                    <h1>Simple html with several headings</h1>
                    <p>Some content 1</p>
                    <h2>Second heading with German vowels ä, ö, and ü</h2>
                    <p>Some content 2</p>
                    <h3>Third</h3>
                    <p>Some content 3</p>
                </body>
            </html>
        """

SOURCE_MARKDOWN_FOR_PPTX = """
# Slide 1 Title

First slide content with bullet points:

- Item 1
- Item 2
- Item 3

---

# Slide 2 Title

Second slide with more content.

Some **bold** and *italic* text.
"""


class TestParameters(NamedTuple):
    base_url: str
    flush_tmp_file_enabled: bool
    request_session: requests.Session
    container: Container
    __test__ = False


def _stop_and_remove_container(container: Container) -> None:
    """Helper function to stop and remove a single container."""
    try:
        logger.info(f"Stopping container: {container.name}")
        container.stop(timeout=1)
    except docker.errors.APIError as e:
        logger.warning(f"Could not stop container {container.name}: {e}")

    try:
        logger.info(f"Removing container: {container.name}")
        container.remove(force=True)
    except docker.errors.APIError as e:
        logger.error(f"Error removing container {container.name}: {e}")


def _remove_image(image) -> None:
    """Helper function to remove a single image."""
    try:
        logger.info(f"Removing image: {image.tags}")
        image.remove(force=True)
    except docker.errors.APIError as e:
        logger.error(f"Error removing image {image.tags}: {e}")


def _is_test_related_container(container: Container) -> bool:
    """Check if a container is related to our tests."""
    return (
        container.name == TEST_CONTAINER_NAME
        or (container.image.tags and TEST_IMAGE_FULL in str(container.image.tags))
        or not container.image.tags  # Intermediate containers
        or (container.image.tags and "debian:trixie-slim" in str(container.image.tags))  # Base image containers
    )


def _is_test_related_image(image) -> bool:
    """Check if an image is related to our tests."""
    return (image.tags and TEST_IMAGE_FULL in str(image.tags)) or not image.tags


def _cleanup_containers(client: docker.DockerClient) -> None:
    """Clean up test-related containers."""
    try:
        containers = client.containers.list(all=True)
        for container in containers:
            if _is_test_related_container(container):
                _stop_and_remove_container(container)
    except docker.errors.APIError as e:
        logger.error(f"Error listing containers: {e}")


def _cleanup_images(client: docker.DockerClient) -> None:
    """Clean up test-related images."""
    try:
        images = client.images.list(all=True)
        for image in images:
            if _is_test_related_image(image):
                _remove_image(image)
    except docker.errors.APIError as e:
        logger.error(f"Error listing images: {e}")


def _verify_containers(client: docker.DockerClient) -> None:
    """Verify and clean up any remaining test-related containers."""
    try:
        remaining = client.containers.list(all=True)
        remaining_test = [c for c in remaining if _is_test_related_container(c)]

        if remaining_test:
            logger.warning(f"Found {len(remaining_test)} test-related containers still remaining after cleanup")
            for container in remaining_test:
                logger.warning(f"Remaining container: {container.name} ({container.id})")
                _stop_and_remove_container(container)
    # Verification reports rather than raises.
    except Exception as e:  # noqa: BLE001
        logger.error(f"Error in container verification: {e}")


def _verify_images(client: docker.DockerClient) -> None:
    """Verify and clean up any remaining test-related images."""
    try:
        remaining_images = client.images.list(all=True)
        remaining_test_images = [i for i in remaining_images if _is_test_related_image(i)]

        if remaining_test_images:
            logger.warning(f"Found {len(remaining_test_images)} test-related images still remaining after cleanup")
            for image in remaining_test_images:
                logger.warning(f"Remaining image: {image.id} (tags: {image.tags})")
                _remove_image(image)
    # Same, for the image check.
    except Exception as e:  # noqa: BLE001
        logger.error(f"Error in image verification: {e}")


def cleanup_docker_resources():
    """
    Cleanup function to remove any leftover test containers and images.
    Ensures thorough cleanup of all test-related Docker resources.
    """
    client = docker.from_env()

    # Initial cleanup
    _cleanup_containers(client)
    _cleanup_images(client)

    # Final verification
    _verify_containers(client)
    _verify_images(client)


def wait_for_container_ready(container: Container, max_wait_time: int = 60) -> None:
    """
    Wait for container to become ready by checking the /version endpoint.

    Args:
        container: Docker container to wait for
        max_wait_time: Maximum time to wait in seconds (default: 60)

    Raises:
        TimeoutError: If container does not become ready within max_wait_time
    """
    start_time = time.time()
    base_url = "http://localhost:9082"

    while time.time() - start_time < max_wait_time:
        try:
            response = requests.get(f"{base_url}/version", timeout=2)
            if response.status_code == 200:
                logger.info("Container is ready")
                return
        except requests.exceptions.RequestException as e:
            logger.debug(f"Container not ready yet, retrying: {e}")
        time.sleep(1)

    # Timeout reached, print logs for debugging
    logs = container.logs().decode("utf-8")
    raise TimeoutError(f"Container did not become ready within {max_wait_time} seconds. Logs:\n{logs}")


def test_container_logs(test_parameters: TestParameters) -> None:
    logs = test_parameters.container.logs()

    assert b"Pandoc service listening port: 9082\n" in logs


def test_convert_html_to_md(test_parameters: TestParameters) -> None:
    expected_content = __load_test_file("tests/data/expected-html-to-md.md")
    response = __send_request(base_url=test_parameters.base_url, request_session=test_parameters.request_session, source_format="html", target_format="markdown", data=SOURCE_HTML)
    assert response.status_code == 200
    assert response.content.decode("utf-8") == expected_content


def test_convert_html_to_textile(test_parameters: TestParameters) -> None:
    expected_content = __load_test_file("tests/data/expected-html-to-textile.textile")
    response = __send_request(base_url=test_parameters.base_url, request_session=test_parameters.request_session, source_format="html", target_format="textile", data=SOURCE_HTML)
    assert response.status_code == 200
    assert response.content.decode("utf-8") == expected_content


def test_convert_html_to_plain(test_parameters: TestParameters) -> None:
    expected_content = __load_test_file("tests/data/expected-html-to-txt.txt")
    response = __send_request(base_url=test_parameters.base_url, request_session=test_parameters.request_session, source_format="html", target_format="plain", data=SOURCE_HTML)
    assert response.status_code == 200
    assert response.content.decode("utf-8") == expected_content


def test_convert_docx_to_plain(test_parameters: TestParameters) -> None:
    with Path("tests/data/test-input.docx").open("rb") as source_file:
        expected_content = __load_test_file("tests/data/expected-docx-to-txt.txt")
        data = ("test-input.docx", source_file.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response = __send_request(base_url=test_parameters.base_url, request_session=test_parameters.request_session, source_format="docx", target_format="plain", data=data)
        assert response.status_code == 200
        assert response.content.decode("utf-8") == expected_content


def test_convert_html_to_docx(test_parameters: TestParameters) -> None:
    response = __send_request(base_url=test_parameters.base_url, request_session=test_parameters.request_session, source_format="html", target_format="docx", data=SOURCE_HTML)
    assert response.status_code == 200

    document = Document(io.BytesIO(response.content))

    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    expected_paragraphs = [
        "Simple html with an ordered list",
        "First",
        "Second",
        "Third",
        "Some bold German vowels ä, ö, and ü at the bottom.",
    ]

    assert expected_paragraphs == paragraphs


def test_convert_with_docx_template(test_parameters: TestParameters) -> None:
    # First test without template - it has some default headings color
    response = __send_docx_with_template_request(base_url=test_parameters.base_url, request_session=test_parameters.request_session, data=SOURCE_HTML_WITH_HEADINGS, source_format="html")
    __assert_doc_contains_specific_headers_color(RGBColor(15, 71, 97), response.content)

    # Now test with 'RED' template - it forces red color for headings
    with Path("tests/data/template-red.docx").open("rb") as t:
        template = t.read()
    response = __send_docx_with_template_request(base_url=test_parameters.base_url, request_session=test_parameters.request_session, data=SOURCE_HTML_WITH_HEADINGS, source_format="html", template=template)
    __assert_doc_contains_specific_headers_color(RGBColor(255, 0, 0), response.content)


def test_version_endpoint(test_parameters: TestParameters) -> None:
    """Test that the /version endpoint returns the expected information."""
    url = f"{test_parameters.base_url}/version"
    response = test_parameters.request_session.get(url)

    # Verify response status
    assert response.status_code == 200

    # Parse response as JSON
    version_info = response.json()

    # Verify all expected fields are present
    assert "apiVersion" in version_info
    assert "python" in version_info
    assert "pandoc" in version_info
    assert "pandocService" in version_info
    assert "timestamp" in version_info
    assert "chromium" in version_info

    # API version is an integer contract version used by docx-exporter
    assert isinstance(version_info["apiVersion"], int)

    # Verify that values are reasonable (not empty where required)
    assert version_info["python"], "Python version should not be empty"
    assert version_info["pandoc"], "Pandoc version should not be empty"
    assert version_info["pandocService"], "Pandoc service version should not be empty"
    assert version_info["chromium"], "Chromium version should not be empty"


def __send_request(base_url: str, request_session: requests.Session, source_format: str, target_format: str, data) -> requests.Response:
    url = f"{base_url}/convert/{source_format}/to/{target_format}"
    files = None
    payload = None

    if isinstance(data, tuple):
        filename, file_content, content_type = data
        files = {"source": (filename, file_content, content_type)}
    else:
        payload = data
    try:
        response = request_session.request(method="POST", url=url, data=payload, files=files, verify=True)
        if response.status_code // 100 != 2:
            logger.error(f"Error: Unexpected response: '{response}'")
            logger.error(f"Error: Response content: '{response.content}'")
    except requests.exceptions.RequestException as e:
        logger.error(f"Error: {e}")
        raise
    else:
        return response


def __send_docx_with_template_request(base_url: str, request_session: requests.Session, source_format: str, data, template=None) -> requests.Response:
    url = f"{base_url}/convert/{source_format}/to/docx-with-template"
    files = {"source": ("file.html", data)}
    if template:
        files["template"] = ("template.docx", template)
    try:
        response = request_session.request(method="POST", url=url, files=files, verify=True)
        if response.status_code // 100 != 2:
            logger.error(f"Error: Unexpected response: '{response}'")
            logger.error(f"Error: Response content: '{response.content}'")
    except requests.exceptions.RequestException as e:
        logger.error(f"Error: {e}")
        raise
    else:
        return response


def __assert_doc_contains_specific_headers_color(color, doc_content):
    document = Document(io.BytesIO(doc_content))

    # Check for specific headings colors and extract their text
    headings = []
    for paragraph in document.paragraphs:
        if paragraph.style.style_id.startswith("Heading"):
            assert color in {paragraph.style.base_style.font.color.rgb, paragraph.style.font.color.rgb}
            headings.append(paragraph.text.replace("\xa0", " "))

    expected_headings = [
        "Simple html with several headings",
        "Second heading with German vowels ä, ö, and ü",
        "Third",
    ]
    assert expected_headings == headings


def __load_test_file(file_path: str) -> str:
    with Path(file_path).open(encoding="utf-8") as file:
        return file.read()


def __send_pptx_with_template_request(base_url: str, request_session: requests.Session, source_format: str, data, template=None) -> requests.Response:
    url = f"{base_url}/convert/{source_format}/to/pptx-with-template"
    files = {"source": ("file.md", data)}
    if template:
        files["template"] = ("template.pptx", template)
    try:
        response = request_session.request(method="POST", url=url, files=files, verify=True)
        if response.status_code // 100 != 2:
            logger.error(f"Error: Unexpected response: '{response}'")
            logger.error(f"Error: Response content: '{response.content}'")
    except requests.exceptions.RequestException as e:
        logger.error(f"Error: {e}")
        raise
    else:
        return response


def test_container_no_error_logs(test_parameters: TestParameters) -> None:
    """Verify container logs contain expected startup messages and no errors."""
    logs = test_parameters.container.logs().decode("utf-8")
    log_lines = logs.splitlines()

    # Check for critical errors (should not contain ERROR or CRITICAL level messages)
    errors = [line for line in log_lines if " - ERROR - " in line or " - CRITICAL - " in line]
    assert not errors, f"Found error logs: {errors}"

    # Check for expected startup message
    assert any("Pandoc service listening port: 9082" in line for line in log_lines), "Expected startup message not found in logs"


def test_docx_template_endpoint(test_parameters: TestParameters) -> None:
    """Test that the /docx-template endpoint returns a valid DOCX file."""
    url = f"{test_parameters.base_url}/docx-template"
    response = test_parameters.request_session.get(url)

    assert response.status_code == 200
    assert response.headers.get("Content-Type") == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert response.headers.get("Content-Disposition") == "attachment; filename=reference.docx"

    # Verify it's a valid DOCX by loading it
    document = Document(io.BytesIO(response.content))
    assert document is not None


def test_pptx_template_endpoint(test_parameters: TestParameters) -> None:
    """Test that the /pptx-template endpoint returns a valid PPTX file."""
    url = f"{test_parameters.base_url}/pptx-template"
    response = test_parameters.request_session.get(url)

    assert response.status_code == 200
    assert response.headers.get("Content-Type") == "application/vnd.openxmlformats-officedocument.presentationml.presentation"
    assert response.headers.get("Content-Disposition") == "attachment; filename=reference.pptx"

    # Verify it's a valid PPTX by loading it and parsing its information
    presentation = io.BytesIO(response.content)
    assert presentation is not None and presentation.getbuffer().nbytes > 0
    find_presentation_information(presentation)


def test_convert_markdown_to_pptx(test_parameters: TestParameters) -> None:
    """Test converting markdown to PPTX format."""
    response = __send_request(
        base_url=test_parameters.base_url,
        request_session=test_parameters.request_session,
        source_format="markdown",
        target_format="pptx",
        data=SOURCE_MARKDOWN_FOR_PPTX,
    )
    assert response.status_code == 200

    # Verify it's a valid PPTX with expected slide count

    presentation = io.BytesIO(response.content)
    assert presentation is not None and presentation.getbuffer().nbytes > 0
    _, _, num_slides = find_presentation_information(presentation)
    assert num_slides == 2, "Expected 2 slides from markdown with slide separator"

    # Check response headers
    assert "Pandoc-Version" in response.headers
    assert "Python-Version" in response.headers


def test_convert_with_pptx_template(test_parameters: TestParameters) -> None:
    """Test PPTX conversion with custom template."""
    # First, get the default template from the endpoint
    template_url = f"{test_parameters.base_url}/pptx-template"
    template_response = test_parameters.request_session.get(template_url)
    assert template_response.status_code == 200
    template = template_response.content

    # Test with template
    response = __send_pptx_with_template_request(
        base_url=test_parameters.base_url,
        request_session=test_parameters.request_session,
        source_format="markdown",
        data=SOURCE_MARKDOWN_FOR_PPTX,
        template=template,
    )
    assert response.status_code == 200

    # Verify it's a valid PPTX
    presentation = io.BytesIO(response.content)
    assert presentation is not None and presentation.getbuffer().nbytes > 0
    _, _, num_slides = find_presentation_information(presentation)
    assert num_slides >= 1


def test_convert_invalid_source_format(test_parameters: TestParameters) -> None:
    """Test that invalid source format returns proper error."""
    response = __send_request(
        base_url=test_parameters.base_url,
        request_session=test_parameters.request_session,
        source_format="invalid_format",
        target_format="html",
        data="test content",
    )
    assert response.status_code == 400


def test_convert_invalid_target_format(test_parameters: TestParameters) -> None:
    """Test that invalid target format returns proper error."""
    response = __send_request(
        base_url=test_parameters.base_url,
        request_session=test_parameters.request_session,
        source_format="markdown",
        target_format="invalid_format",
        data="test content",
    )
    assert response.status_code == 400


SOURCE_HTML_WITH_LOCAL_FILE_REFERENCES = """
            <html>
                <body>
                    <p>A document naming resources of the container it is converted in.</p>
                    <img src="/etc/hostname"/>
                    <img src="file:///etc/hostname"/>
                    <img src="/etc/passwd"/>
                </body>
            </html>
            """


def test_a_document_cannot_read_files_of_the_container(test_parameters: TestParameters) -> None:
    """The sandbox keeps a document from naming a path of this container.

    Without it, pandoc reads the file and embeds it in the result: an exfiltration
    channel out of every deployment reachable by a caller.
    """
    response = __send_request(
        base_url=test_parameters.base_url,
        request_session=test_parameters.request_session,
        source_format="html",
        target_format="docx",
        data=SOURCE_HTML_WITH_LOCAL_FILE_REFERENCES,
    )
    assert response.status_code == 200

    with zipfile.ZipFile(io.BytesIO(response.content)) as docx:
        embedded = [name for name in docx.namelist() if name.startswith("word/media/")]
        assert embedded == [], f"the document pulled files of the container into the result: {embedded}"
        document_xml = docx.read("word/document.xml").decode("utf-8")

    assert "root:x:0:0" not in document_xml


SOURCE_MARKDOWN_WITH_RAW_TEX = """Hello

\\input{/etc/hostname}
"""


def test_a_document_cannot_reach_the_pdf_engine_with_raw_tex(test_parameters: TestParameters) -> None:
    """The PDF is produced by tectonic, which runs outside the pandoc sandbox.

    Raw TeX of the document is therefore dropped before it gets there, or
    ``\\input{/etc/hostname}`` would put a file of the container into the PDF.
    The document with the raw TeX has to render exactly like the one without it.
    """

    def text_of(markdown: str) -> str:
        response = __send_request(
            base_url=test_parameters.base_url,
            request_session=test_parameters.request_session,
            source_format="markdown",
            target_format="pdf",
            data=markdown,
        )
        assert response.status_code == 200
        reader = PdfReader(io.BytesIO(response.content))
        return "".join(page.extract_text() for page in reader.pages)

    with_raw_tex = text_of(SOURCE_MARKDOWN_WITH_RAW_TEX)
    without_raw_tex = text_of("Hello\n")

    assert "Hello" in without_raw_tex
    assert with_raw_tex == without_raw_tex


SOURCE_MARKDOWN_WITH_TEX_IN_MATH = """Hello

$\\input{/etc/hostname}$
"""

SOURCE_MARKDOWN_WITH_A_LOCAL_IMAGE = """Hello

![](/opt/pandoc/.build_timestamp)
"""


def _pdf_text_and_images(test_parameters: TestParameters, markdown: str) -> tuple[str, int]:
    """Convert a markdown source to PDF and report what the result carries."""
    response = __send_request(
        base_url=test_parameters.base_url,
        request_session=test_parameters.request_session,
        source_format="markdown",
        target_format="pdf",
        data=markdown,
    )
    assert response.status_code == 200
    reader = PdfReader(io.BytesIO(response.content))
    text = "".join(page.extract_text() for page in reader.pages)
    return text, sum(_image_count(page) for page in reader.pages)


def _image_count(page: object) -> int:
    """Count the images of a page by its resources, which needs no image library."""
    resources = page.get("/Resources")  # type: ignore[attr-defined]
    if resources is None:
        return 0
    xobjects = resources.get_object().get("/XObject")
    if xobjects is None:
        return 0
    return sum(1 for entry in xobjects.get_object().values() if entry.get_object().get("/Subtype") == "/Image")


def test_math_cannot_carry_tex_to_the_pdf_engine(test_parameters: TestParameters) -> None:
    """The writer emits math verbatim, so the same primitive inside $...$ takes the same route."""
    with_tex, _ = _pdf_text_and_images(test_parameters, SOURCE_MARKDOWN_WITH_TEX_IN_MATH)
    plain, _ = _pdf_text_and_images(test_parameters, "Hello\n")

    assert with_tex == plain


def test_a_document_cannot_put_a_file_of_the_container_into_a_pdf_as_an_image(test_parameters: TestParameters) -> None:
    """An image path becomes \\includegraphics, which the engine outside the sandbox resolves."""
    _, images = _pdf_text_and_images(test_parameters, SOURCE_MARKDOWN_WITH_A_LOCAL_IMAGE)

    assert images == 0


SOURCE_MARKDOWN_WITH_MIXED_CASE_RAW_TEX = """Hello

```{=LaTeX}
\\input{/etc/hostname}
```
"""

SOURCE_MARKDOWN_WITH_AN_EMBEDDED_IMAGE = "Hello\n\n![](data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAgAAAAIAQMAAAD+wSzIAAAABlBMVEX///+/v7+jQ3Y5AAAADklEQVQI12P4AIX8EAgALgAD/aNpbtEAAAAASUVORK5CYII=)\n"


def test_the_raw_tex_check_is_case_insensitive(test_parameters: TestParameters) -> None:
    """Pandoc folds the case of a raw format, so `{=LaTeX}` takes the same route as `{=latex}`."""
    with_raw_tex, _ = _pdf_text_and_images(test_parameters, SOURCE_MARKDOWN_WITH_MIXED_CASE_RAW_TEX)
    plain, _ = _pdf_text_and_images(test_parameters, "Hello\n")

    assert with_raw_tex == plain


def test_an_image_carried_by_the_document_reaches_the_pdf(test_parameters: TestParameters) -> None:
    """The positive side of the rule: what travels inside the document is kept."""
    _, images = _pdf_text_and_images(test_parameters, SOURCE_MARKDOWN_WITH_AN_EMBEDDED_IMAGE)

    assert images == 1


MATH_WHICH_REACHES_FOR_A_FILE = {
    "carets": "$^^5cinput{/etc/hostname}$",
    "wide carets": "$^^^^005cinput{/etc/hostname}$",
    "kernel name": "$\\makeatletter\\@@input{/etc/hostname}$",
    "other kernel name": "$\\makeatletter\\@input{/etc/hostname}$",
    "alias by def": "$\\def\\x{\\input}\\x{/etc/hostname}$",
    "alias by let": "$\\let\\x\\input \\x{/etc/hostname}$",
    "name built with csname": "$\\csname input\\endcsname{/etc/hostname}$",
    "name built with scantokens": "$\\scantokens{\\string\\i nput{/etc/hostname}}$",
    "a file read as a font": "$\\font\\x=/etc/hostname \\x$",
    "a file embedded by the engine": '$\\XeTeXpdffile"/etc/hostname"$',
}

ORDINARY_MATH = {
    "a power": "$E = mc^2$",
    "a fraction, a greek letter and a sum": "$\\frac{a}{b} + \\alpha \\sum_{i=1}^{n} x_i$",
    "a root and an integral": "$\\sqrt{x} \\int_0^\\infty e^{-t}\\,dt$",
    "a matrix": "$\\begin{pmatrix} 1 & 0 \\\\ 0 & 1 \\end{pmatrix}$",
    "text inside math": "$x \\text{ for all } x$",
}


@pytest.mark.parametrize("math", MATH_WHICH_REACHES_FOR_A_FILE.values(), ids=MATH_WHICH_REACHES_FOR_A_FILE.keys())
def test_math_cannot_reach_a_file_however_it_spells_it(test_parameters: TestParameters, math: str) -> None:
    """A formula may not name a primitive, build a name, or carry the `@` of the kernel."""
    reaching, _ = _pdf_text_and_images(test_parameters, f"Hello\n\n{math}\n")
    plain, _ = _pdf_text_and_images(test_parameters, "Hello\n")

    assert reaching == plain


@pytest.mark.parametrize("math", ORDINARY_MATH.values(), ids=ORDINARY_MATH.keys())
def test_ordinary_math_still_reaches_the_pdf(test_parameters: TestParameters, math: str) -> None:
    """The positive side of the rule: a formula which reaches for nothing renders as before."""
    rendered, _ = _pdf_text_and_images(test_parameters, f"Hello\n\n{math}\n")
    plain, _ = _pdf_text_and_images(test_parameters, "Hello\n")

    assert rendered != plain
