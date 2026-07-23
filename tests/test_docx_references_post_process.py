import logging
import re
from unittest.mock import MagicMock

from docx.document import Document as DocumentObject
from docx.oxml import parse_xml
from lxml import etree

from app.DocxReferencesPostProcess import (
    KIND_TOC,
    KIND_TOF,
    KIND_TOT,
    SCHEMA,
    TOF_FIELD_CODE,
    _create_field_with_entries,
    _get_max_bookmark_id,
    _get_seq_name,
    _parse_placeholder,
    _add_tc_field,
    _create_toc_field,
    _create_tof_field,
    _create_tot_field,
    _set_paragraph_style,
    _ensure_seq_field,
    _find_placeholder_paragraphs,
    _get_paragraph_style,
    _get_paragraph_text,
    _has_seq_field,
    _is_adjacent_to_table,
    add_table_of_contents_entries,
    enable_auto_update_fields,
)

_MINIMAL_TBL = f'<w:tbl xmlns:w="{SCHEMA}"><w:tr><w:tc><w:p><w:r><w:t>x</w:t></w:r></w:p></w:tc></w:tr></w:tbl>'


# ---- _get_paragraph_text / _get_paragraph_style ----


def test_get_paragraph_text_empty():
    assert _get_paragraph_text(parse_xml(f'<w:p xmlns:w="{SCHEMA}"/>')) == ""


def test_get_paragraph_text_single():
    assert _get_paragraph_text(parse_xml(f'<w:p xmlns:w="{SCHEMA}"><w:r><w:t>Hi</w:t></w:r></w:p>')) == "Hi"


def test_get_paragraph_style_none():
    assert _get_paragraph_style(parse_xml(f'<w:p xmlns:w="{SCHEMA}"/>')) is None


def test_get_paragraph_style_bodytext():
    assert _get_paragraph_style(parse_xml(f'<w:p xmlns:w="{SCHEMA}"><w:pPr><w:pStyle w:val="BodyText"/></w:pPr></w:p>')) == "BodyText"


# ---- _set_paragraph_style ----


def test_set_paragraph_style_adds():
    para = parse_xml(f'<w:p xmlns:w="{SCHEMA}"><w:r><w:t>x</w:t></w:r></w:p>')
    _set_paragraph_style(para, "Caption")
    assert para.find(".//w:pStyle", namespaces={"w": SCHEMA}).get(f"{{{SCHEMA}}}val") == "Caption"


def test_set_paragraph_style_replaces():
    para = parse_xml(f'<w:p xmlns:w="{SCHEMA}"><w:pPr><w:pStyle w:val="Normal"/></w:pPr></w:p>')
    _set_paragraph_style(para, "Caption")
    assert para.find(".//w:pStyle", namespaces={"w": SCHEMA}).get(f"{{{SCHEMA}}}val") == "Caption"


# ---- _is_adjacent_to_table ----


def test_adjacent_to_table_before():
    body = parse_xml(f'<w:body xmlns:w="{SCHEMA}"><w:p><w:r><w:t>c</w:t></w:r></w:p>{_MINIMAL_TBL}</w:body>')
    assert _is_adjacent_to_table(body.find("w:p", namespaces={"w": SCHEMA})) is True


def test_after_table_is_not_adjacent():
    """A paragraph after a table should NOT be classified as table caption."""
    body = parse_xml(f'<w:body xmlns:w="{SCHEMA}">{_MINIMAL_TBL}<w:p><w:r><w:t>c</w:t></w:r></w:p></w:body>')
    paras = body.findall("w:p", namespaces={"w": SCHEMA})
    assert _is_adjacent_to_table(paras[-1]) is False


def test_adjacent_to_table_skips_bookmarks():
    body = parse_xml(f'''<w:body xmlns:w="{SCHEMA}">
        <w:p><w:r><w:t>c</w:t></w:r></w:p>
        <w:bookmarkStart w:id="1" w:name="x"/>
        {_MINIMAL_TBL}
        <w:bookmarkEnd w:id="1"/>
    </w:body>''')
    assert _is_adjacent_to_table(body.find("w:p", namespaces={"w": SCHEMA})) is True


def test_adjacent_to_table_skips_empty_paras():
    body = parse_xml(f'<w:body xmlns:w="{SCHEMA}"><w:p><w:r><w:t>c</w:t></w:r></w:p><w:p/>{_MINIMAL_TBL}</w:body>')
    assert _is_adjacent_to_table(body.find("w:p", namespaces={"w": SCHEMA})) is True


def test_not_adjacent_to_table():
    body = parse_xml(f'<w:body xmlns:w="{SCHEMA}"><w:p><w:r><w:t>a</w:t></w:r></w:p><w:p><w:r><w:t>b</w:t></w:r></w:p></w:body>')
    assert _is_adjacent_to_table(body.find("w:p", namespaces={"w": SCHEMA})) is False


# ---- _has_seq_field / _ensure_seq_field ----


def test_has_seq_field_true():
    para = parse_xml(f'<w:p xmlns:w="{SCHEMA}"><w:r><w:instrText> SEQ Table </w:instrText></w:r></w:p>')
    assert _has_seq_field(para) is True


def test_has_seq_field_false():
    para = parse_xml(f'<w:p xmlns:w="{SCHEMA}"><w:r><w:t>Table 1</w:t></w:r></w:p>')
    assert _has_seq_field(para) is False


def test_ensure_seq_field_replaces_plain_number():
    para = parse_xml(f'<w:p xmlns:w="{SCHEMA}"><w:pPr><w:pStyle w:val="Caption"/></w:pPr><w:r><w:t>Figure 1 Pic</w:t></w:r></w:p>')
    _ensure_seq_field(para, "Figure")
    xml = etree.tostring(para, encoding="unicode")
    assert "SEQ Figure" in xml
    assert 'fldCharType="begin"' in xml
    assert re.search(r"separate.*<w:t[^>]*>1</w:t>.*end", xml, re.S)
    assert "Figure " in xml
    assert " Pic" in xml


def test_ensure_seq_field_skips_existing():
    para = parse_xml(f'''<w:p xmlns:w="{SCHEMA}">
        <w:r><w:fldChar w:fldCharType="begin"/></w:r>
        <w:r><w:instrText> SEQ Table \\* ARABIC </w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="end"/></w:r>
    </w:p>''')
    xml_before = etree.tostring(para, encoding="unicode")
    _ensure_seq_field(para, "Table")
    assert etree.tostring(para, encoding="unicode") == xml_before


def test_ensure_seq_field_no_number():
    para = parse_xml(f'<w:p xmlns:w="{SCHEMA}"><w:r><w:t>No number</w:t></w:r></w:p>')
    xml_before = etree.tostring(para, encoding="unicode")
    _ensure_seq_field(para, "Table")
    assert etree.tostring(para, encoding="unicode") == xml_before


# ---- _add_tc_field ----


def test_add_tc_field_creates_field_with_bookmark():
    para = parse_xml(f'<w:p xmlns:w="{SCHEMA}"><w:r><w:t>Table 1</w:t></w:r></w:p>')
    _add_tc_field(para, "Table 1", "T", 42, "_Toc000000042")
    xml = etree.tostring(para, encoding="unicode")
    assert "bookmarkStart" in xml
    assert "_Toc000000042" in xml
    assert "\\f T" in xml


# ---- TOC/TOF/TOT field creation ----


def test_create_toc_field():
    paras = _create_toc_field()
    assert len(paras) == 2
    instr = paras[0].find(".//w:instrText", namespaces={"w": SCHEMA})
    assert '\\o "1-9"' in instr.text


def test_create_tof_field():
    paras = _create_tof_field()
    instr = paras[0].find(".//w:instrText", namespaces={"w": SCHEMA})
    assert "\\f F" in instr.text


def test_create_tot_field():
    paras = _create_tot_field()
    instr = paras[0].find(".//w:instrText", namespaces={"w": SCHEMA})
    assert "\\f T" in instr.text


def test_tot_with_entries_has_hyperlinks():
    entries = [("Table 1", "_Toc1"), ("Table 2", "_Toc2")]
    xml = "".join(etree.tostring(p, encoding="unicode") for p in _create_tot_field(entries))
    assert 'w:anchor="_Toc1"' in xml
    assert 'w:anchor="_Toc2"' in xml
    assert "PAGEREF" in xml


def test_tof_with_entries():
    entries = [("Figure 1", "_Toc10")]
    xml = "".join(etree.tostring(p, encoding="unicode") for p in _create_tof_field(entries))
    assert 'w:anchor="_Toc10"' in xml
    assert "\\f F" in xml


# ---- placeholder finding ----


def test_find_placeholders_all_three():
    body = parse_xml(f'''<w:body xmlns:w="{SCHEMA}">
        <w:p><w:r><w:t>TOC_PLACEHOLDER</w:t></w:r></w:p>
        <w:p><w:r><w:t>TOF_PLACEHOLDER</w:t></w:r></w:p>
        <w:p><w:r><w:t>TOT_PLACEHOLDER</w:t></w:r></w:p>
    </w:body>''')
    result = []
    _find_placeholder_paragraphs(body, result)
    assert len(result) == 3


def test_find_placeholders_ignores_heading():
    body = parse_xml(f'<w:body xmlns:w="{SCHEMA}"><w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr><w:r><w:t>TOC_PLACEHOLDER</w:t></w:r></w:p></w:body>')
    result = []
    _find_placeholder_paragraphs(body, result)
    assert len(result) == 0


# ---- enable_auto_update_fields ----


def test_enable_auto_update_fields():
    mock_doc = MagicMock(spec=DocumentObject)
    mock_update_fields = MagicMock()
    mock_doc.settings.element.find.return_value = mock_update_fields
    enable_auto_update_fields(mock_doc)
    mock_update_fields.set.assert_called_once()


def test_enable_auto_update_fields_exception(caplog):
    mock_doc = MagicMock(spec=DocumentObject)
    mock_doc.settings.element.find.side_effect = Exception("err")
    with caplog.at_level(logging.WARNING):
        enable_auto_update_fields(mock_doc)
    assert "Could not enable auto-update fields" in caplog.text


# ---- Integration: add_table_of_contents_entries ----


def test_figure_captions_found(caplog):
    mock_doc = MagicMock(spec=DocumentObject)
    body = parse_xml(f'<w:body xmlns:w="{SCHEMA}"><w:p><w:pPr><w:pStyle w:val="Caption"/></w:pPr><w:r><w:t>Figure 1</w:t></w:r></w:p></w:body>')
    mock_doc.element.body = body
    with caplog.at_level(logging.INFO):
        add_table_of_contents_entries(mock_doc)
    assert "Found 1 figure captions" in caplog.text


def test_table_captions_found(caplog):
    mock_doc = MagicMock(spec=DocumentObject)
    body = parse_xml(f'<w:body xmlns:w="{SCHEMA}"><w:p><w:pPr><w:pStyle w:val="Caption"/></w:pPr><w:r><w:t>Table 1</w:t></w:r></w:p>{_MINIMAL_TBL}</w:body>')
    mock_doc.element.body = body
    with caplog.at_level(logging.INFO):
        add_table_of_contents_entries(mock_doc)
    assert "1 table captions" in caplog.text


def test_non_caption_paragraphs_ignored(caplog):
    mock_doc = MagicMock(spec=DocumentObject)
    body = parse_xml(f'<w:body xmlns:w="{SCHEMA}"><w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr><w:r><w:t>Table 1</w:t></w:r></w:p></w:body>')
    mock_doc.element.body = body
    with caplog.at_level(logging.INFO):
        add_table_of_contents_entries(mock_doc)
    assert "Found 0 figure captions and 0 table captions" in caplog.text


def test_captions_get_unique_bookmarks():
    mock_doc = MagicMock(spec=DocumentObject)
    body = parse_xml(f'''<w:body xmlns:w="{SCHEMA}">
        <w:p><w:pPr><w:pStyle w:val="Caption"/></w:pPr><w:r><w:t>Table 1</w:t></w:r></w:p>{_MINIMAL_TBL}
        <w:p><w:pPr><w:pStyle w:val="Caption"/></w:pPr><w:r><w:t>Table 2</w:t></w:r></w:p>{_MINIMAL_TBL}
    </w:body>''')
    mock_doc.element.body = body
    add_table_of_contents_entries(mock_doc)
    xml = etree.tostring(body, encoding="unicode")
    bookmarks = re.findall(r'w:name="(_Toc\d+)"', xml)
    assert len(bookmarks) >= 2
    assert bookmarks[0] != bookmarks[1]


def test_tot_placeholder_full_workflow():
    mock_doc = MagicMock(spec=DocumentObject)
    body = parse_xml(f'''<w:body xmlns:w="{SCHEMA}">
        <w:p><w:pPr><w:pStyle w:val="Caption"/></w:pPr><w:r><w:t>Table 1</w:t></w:r></w:p>{_MINIMAL_TBL}
        <w:p><w:r><w:t>TOT_PLACEHOLDER</w:t></w:r></w:p>
    </w:body>''')
    mock_doc.element.body = body
    add_table_of_contents_entries(mock_doc)
    xml = etree.tostring(body, encoding="unicode")
    assert "bookmarkStart" in xml
    assert "w:hyperlink" in xml
    assert "PAGEREF" in xml
    assert "\\f T" in xml


def test_localized_caption_classified_by_table_adjacency():
    """Polish 'Tabela 1' next to a table should be classified as table caption."""
    mock_doc = MagicMock(spec=DocumentObject)
    body = parse_xml(f'''<w:body xmlns:w="{SCHEMA}">
        <w:p><w:pPr><w:pStyle w:val="Caption"/></w:pPr>
            <w:r><w:t xml:space="preserve">Tabela </w:t></w:r>
            <w:r><w:fldChar w:fldCharType="begin"/></w:r>
            <w:r><w:instrText xml:space="preserve"> SEQ Tabela \\* ARABIC </w:instrText></w:r>
            <w:r><w:fldChar w:fldCharType="separate"/></w:r>
            <w:r><w:t>1</w:t></w:r>
            <w:r><w:fldChar w:fldCharType="end"/></w:r>
        </w:p>{_MINIMAL_TBL}
    </w:body>''')
    mock_doc.element.body = body
    add_table_of_contents_entries(mock_doc)
    xml = etree.tostring(body, encoding="unicode")
    assert "\\f T" in xml


def test_full_workflow_empty():
    from docx import Document

    doc = Document()
    add_table_of_contents_entries(doc)


def test_enable_auto_update_real():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Test")
    enable_auto_update_fields(doc)
    uf = doc.settings.element.find(".//w:updateFields", namespaces={"w": SCHEMA})
    assert uf is not None
    assert uf.get(f"{{{SCHEMA}}}val") == "true"


# ---- helper edge cases (coverage-complete) ----


def test_parse_placeholder_variants():
    assert _parse_placeholder("TOC_PLACEHOLDER") == KIND_TOC
    assert _parse_placeholder("TOF_PLACEHOLDER") == KIND_TOF
    assert _parse_placeholder("TOT_PLACEHOLDER") == KIND_TOT
    assert _parse_placeholder("Regular text") is None


def test_get_max_bookmark_id_skips_non_integer_ids():
    body = parse_xml(f'''<w:body xmlns:w="{SCHEMA}">
        <w:p><w:bookmarkStart w:id="5" w:name="a"/><w:bookmarkEnd w:id="5"/></w:p>
        <w:p><w:bookmarkStart w:id="abc" w:name="b"/><w:bookmarkEnd w:id="abc"/></w:p>
    </w:body>''')
    assert _get_max_bookmark_id(body) == 5


def test_image_caption_style_is_classified_as_figure():
    mock_doc = MagicMock(spec=DocumentObject)
    body = parse_xml(f'''<w:body xmlns:w="{SCHEMA}">
        <w:p><w:pPr><w:pStyle w:val="ImageCaption"/></w:pPr><w:r><w:t>Abbildung 1 Bild</w:t></w:r></w:p>
    </w:body>''')
    mock_doc.element.body = body
    add_table_of_contents_entries(mock_doc)
    xml = etree.tostring(body, encoding="unicode")
    assert "SEQ Figure" in xml


def test_get_seq_name_reads_identifier():
    para = parse_xml(f'''<w:p xmlns:w="{SCHEMA}">
        <w:r><w:instrText xml:space="preserve"> SEQ Tabelle \\* ARABIC </w:instrText></w:r>
    </w:p>''')
    assert _get_seq_name(para) == "Tabelle"


def test_get_seq_name_without_seq_field():
    para = parse_xml(f'<w:p xmlns:w="{SCHEMA}"><w:r><w:t>Table 1</w:t></w:r></w:p>')
    assert _get_seq_name(para) is None


def test_get_seq_name_with_malformed_instruction():
    # First instruction mentions SEQ without an identifier, second is valid:
    # the scan must skip the malformed one and keep looking
    para = parse_xml(f'''<w:p xmlns:w="{SCHEMA}">
        <w:r><w:instrText xml:space="preserve"> PAGEREF _Toc1 \\h </w:instrText></w:r>
        <w:r><w:instrText xml:space="preserve">SEQ</w:instrText></w:r>
        <w:r><w:instrText xml:space="preserve"> SEQ Tabelle \\* ARABIC </w:instrText></w:r>
    </w:p>''')
    assert _get_seq_name(para) == "Tabelle"


def test_ensure_seq_field_skips_runs_without_text_or_number():
    para = parse_xml(f'''<w:p xmlns:w="{SCHEMA}">
        <w:r></w:r>
        <w:r><w:t>no number here</w:t></w:r>
        <w:r><w:t>Table 7 caption</w:t></w:r>
    </w:p>''')
    _ensure_seq_field(para, "Table")
    xml = etree.tostring(para, encoding="unicode")
    assert "SEQ Table" in xml
    assert "no number here" in _get_paragraph_text(para)


def test_ensure_seq_field_number_at_run_start():
    # No "before" text: run starts with the number itself
    para = parse_xml(f'<w:p xmlns:w="{SCHEMA}"><w:r><w:t>7 caption tail</w:t></w:r></w:p>')
    _ensure_seq_field(para, "Figure")
    text = _get_paragraph_text(para)
    assert text.startswith("7")
    assert "caption tail" in text
    assert "SEQ Figure" in etree.tostring(para, encoding="unicode")


def test_ensure_seq_field_preserves_run_formatting():
    para = parse_xml(f'''<w:p xmlns:w="{SCHEMA}">
        <w:r><w:rPr><w:b/></w:rPr><w:t>Table 3 bold caption</w:t></w:r>
    </w:p>''')
    _ensure_seq_field(para, "Table")
    # The cloned before/after runs keep the original bold rPr
    runs_with_bold = [r for r in para.findall("w:r", namespaces={"w": SCHEMA}) if r.find("w:rPr/w:b", namespaces={"w": SCHEMA}) is not None]
    assert len(runs_with_bold) >= 2


def test_set_paragraph_style_inserts_into_existing_ppr():
    para = parse_xml(f'<w:p xmlns:w="{SCHEMA}"><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')
    _set_paragraph_style(para, "Caption")
    assert para.find(".//w:pStyle", namespaces={"w": SCHEMA}).get(f"{{{SCHEMA}}}val") == "Caption"
    assert para.find(".//w:jc", namespaces={"w": SCHEMA}) is not None


def test_find_placeholders_skips_non_paragraph_elements():
    body = parse_xml(f'''<w:body xmlns:w="{SCHEMA}">
        {_MINIMAL_TBL}
        <w:p><w:r><w:t>TOC_PLACEHOLDER</w:t></w:r></w:p>
    </w:body>''')
    result = []
    _find_placeholder_paragraphs(body, result)
    assert len(result) == 1
    assert result[0][2] == KIND_TOC


def test_full_flow_prefills_tof_and_tot_with_entries():
    mock_doc = MagicMock(spec=DocumentObject)
    body = parse_xml(f'''<w:body xmlns:w="{SCHEMA}">
        <w:p><w:r><w:t>TOF_PLACEHOLDER</w:t></w:r></w:p>
        <w:p><w:r><w:t>TOT_PLACEHOLDER</w:t></w:r></w:p>
        <w:p><w:pPr><w:pStyle w:val="Caption"/></w:pPr><w:r><w:t>Figure 1 A picture</w:t></w:r></w:p>
        <w:p><w:pPr><w:pStyle w:val="Caption"/></w:pPr><w:r><w:t>Table 1 A table</w:t></w:r></w:p>{_MINIMAL_TBL}
    </w:body>''')
    mock_doc.element.body = body
    add_table_of_contents_entries(mock_doc)
    xml = etree.tostring(body, encoding="unicode")
    assert "\\f F" in xml
    assert "\\f T" in xml
    # Pre-filled cached entries hyperlink to the captions via PAGEREF
    assert xml.count("PAGEREF") == 2
    assert "TOF_PLACEHOLDER" not in xml
    assert "TOT_PLACEHOLDER" not in xml


def test_placeholder_without_matching_captions_removed_without_field():
    mock_doc = MagicMock(spec=DocumentObject)
    body = parse_xml(f'''<w:body xmlns:w="{SCHEMA}">
        <w:p><w:r><w:t>TOF_PLACEHOLDER</w:t></w:r></w:p>
    </w:body>''')
    mock_doc.element.body = body
    add_table_of_contents_entries(mock_doc)
    xml = etree.tostring(body, encoding="unicode")
    assert "TOF_PLACEHOLDER" not in xml
    assert "\\f F" not in xml


def test_create_field_with_entries_empty_list_falls_back_to_plain_field():
    paragraphs = _create_field_with_entries(TOF_FIELD_CODE, [])
    assert len(paragraphs) == 2  # plain field + spacing paragraph
    xml = etree.tostring(paragraphs[0], encoding="unicode")
    assert "\\f F" in xml
    assert "PAGEREF" not in xml


def test_get_paragraph_text_skips_empty_t_elements():
    para = parse_xml(f'<w:p xmlns:w="{SCHEMA}"><w:r><w:t/></w:r><w:r><w:t>x</w:t></w:r></w:p>')
    assert _get_paragraph_text(para) == "x"


def test_get_paragraph_style_ppr_without_pstyle():
    para = parse_xml(f'<w:p xmlns:w="{SCHEMA}"><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')
    assert _get_paragraph_style(para) is None


def test_full_flow_inserts_toc_field():
    mock_doc = MagicMock(spec=DocumentObject)
    body = parse_xml(f'''<w:body xmlns:w="{SCHEMA}">
        <w:p><w:r><w:t>TOC_PLACEHOLDER</w:t></w:r></w:p>
        <w:p><w:r><w:t>Ordinary body paragraph</w:t></w:r></w:p>
    </w:body>''')
    mock_doc.element.body = body
    add_table_of_contents_entries(mock_doc)
    xml = etree.tostring(body, encoding="unicode")
    assert 'TOC \\o "1-9" \\h \\z \\u' in xml
    assert "TOC_PLACEHOLDER" not in xml
    assert "Ordinary body paragraph" in xml
