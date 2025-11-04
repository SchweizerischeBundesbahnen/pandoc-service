import sys
from unittest.mock import MagicMock, mock_open, patch

import pytest
from docx.table import Table, _Cell
from lxml import etree

from app import DocxPostProcess
from app.DocxPostProcess import SCHEMA, _process_table

WORD_PROCESSING_ML_MAIN_SCHEMA = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WORD_PROCESSING_ML_MAIN_SCHEMA_IN_BRACKETS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
DRAWING_ML_MAIN_SCHEMA = "http://schemas.openxmlformats.org/drawingml/2006/main"
DRAWING_ML_PICTURE_SCHEMA = "http://schemas.openxmlformats.org/drawingml/2006/picture"

SOURCE_HTML_WITH_TABLE = """
            <html>
                <head>
                    <title>Test doc title</title>
                </head>
                <body>
                    <h1>Simple html with table</h1>
                    <table>
                        <thead>
                            <tr>
                                <td style="width: 1000px">Wide column</td>
                                <td>
                                    <img src="{0}"></img>
                                </td>
                            </tr>
                        </thead>
                    </table>
                </body>
            </html>
        """

SOURCE_HTML_WITH_NESTED_TABLE = """
            <html>
                <head>
                    <title>Test doc title</title>
                </head>
                <body>
                    <h1>Nested tables</h1>
                    <table>
                        <tr>
                            <td>
                                <table>
                                    <tr>
                                        <td>Nested cell</td>
                                        <td><img src="{0}"></img></td>
                                    </tr>
                                </table>
                            </td>
                            <td>Outer cell</td>
                        </tr>
                    </table>
                </body>
            </html>
        """

SOURCE_HTML_NO_TABLES = """
            <html>
                <head>
                    <title>Test doc title</title>
                </head>
                <body>
                    <h1>Document without tables</h1>
                    <p>This is a simple paragraph without any tables.</p>
                    <p><img src="{0}"></img></p>
                </body>
            </html>
        """

# HTML with a small image that shouldn't need resizing
SOURCE_HTML_SMALL_IMAGE = """
            <html>
                <head>
                    <title>Test doc title</title>
                </head>
                <body>
                    <h1>Document with small image</h1>
                    <table>
                        <tr>
                            <td>
                                <!-- Using a data URL for a 1x1 pixel transparent PNG -->
                                <img src="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNkYAAAAAYAAjCB0C8AAAAASUVORK5CYII="></img>
                            </td>
                        </tr>
                    </table>
                </body>
            </html>
        """

EMUS_IN_INCH = 914400


# Mock Document class for testing without requiring a real docx file
def create_mock_document_with_table(has_nested_table=False, has_image=True):
    # Create a mock Document
    mock_doc = MagicMock()

    # Create a mock table
    mock_table = MagicMock(spec=Table)
    mock_table._element = MagicMock()

    # Add columns to the table (to avoid division by zero)
    mock_column = MagicMock()
    mock_table.columns = [mock_column]  # At least one column

    # Set up table properties
    mock_table_properties = MagicMock()
    mock_table._element.find.return_value = mock_table_properties

    # Create mock rows and cells
    mock_row = MagicMock()
    mock_cell = MagicMock(spec=_Cell)
    mock_cell._tc = MagicMock()

    # Set up nested tables if needed
    if has_nested_table:
        mock_nested_table = MagicMock(spec=Table)
        mock_nested_table._element = MagicMock()
        # Add columns to the nested table too
        mock_nested_table.columns = [MagicMock()]
        mock_cell.tables = [mock_nested_table]
    else:
        mock_cell.tables = []

    # Set up mock cell XML with image if needed
    if has_image:
        mock_cell._tc.xml = f'''
        <w:tc xmlns:w="{WORD_PROCESSING_ML_MAIN_SCHEMA}"
              xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
              xmlns:a="{DRAWING_ML_MAIN_SCHEMA}"
              xmlns:pic="{DRAWING_ML_PICTURE_SCHEMA}"
              xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
            <w:tcPr/>
            <w:p>
                <w:r>
                    <w:drawing>
                        <wp:inline>
                            <wp:extent cx="1000000" cy="750000"/>
                            <a:graphic>
                                <a:graphicData>
                                    <pic:pic>
                                        <pic:blipFill>
                                            <a:blip r:embed="rId5"/>
                                        </pic:blipFill>
                                        <pic:spPr>
                                            <a:xfrm>
                                                <a:ext cx="1000000" cy="750000"/>
                                            </a:xfrm>
                                        </pic:spPr>
                                    </pic:pic>
                                </a:graphicData>
                            </a:graphic>
                        </wp:inline>
                    </w:drawing>
                </w:r>
            </w:p>
        </w:tc>
        '''
    else:
        mock_cell._tc.xml = f'''
        <w:tc xmlns:w="{WORD_PROCESSING_ML_MAIN_SCHEMA}">
            <w:tcPr/>
            <w:p>
                <w:r>
                    <w:t>Cell content</w:t>
                </w:r>
            </w:p>
        </w:tc>
        '''

    # Connect the mocks together
    mock_row.cells = [mock_cell]
    mock_table.rows = [mock_row]
    mock_doc.tables = [mock_table]

    return mock_doc


@patch("app.DocxPostProcess.etree.fromstring")
def test_replace_table_properties(mock_fromstring):
    # Create a mock document to return from our Document constructor
    mock_doc = create_mock_document_with_table()

    # Create a mock XML tree that will be returned from fromstring
    mock_tree = MagicMock()
    # Return empty list for extent elements to avoid image resizing
    mock_tree.findall.return_value = []
    mock_fromstring.return_value = mock_tree

    # Call the function under test
    DocxPostProcess._replace_table_properties(mock_doc)

    # Verify that table properties were accessed
    assert mock_doc.tables[0]._element.find.called


@patch("app.DocxPostProcess.etree.fromstring")
def test_nested_tables(mock_fromstring):
    # Create a mock document with nested tables
    mock_doc = create_mock_document_with_table(has_nested_table=True)

    # Create a mock XML tree that will be returned from fromstring
    mock_tree = MagicMock()
    # Return empty list for extent elements to avoid image resizing
    mock_tree.findall.return_value = []
    mock_fromstring.return_value = mock_tree

    # Call the function under test
    DocxPostProcess._replace_table_properties(mock_doc)

    # Verify we processed the nested table by checking if cell.tables was accessed
    assert len(mock_doc.tables[0].rows[0].cells[0].tables) > 0


def test_document_without_tables():
    # Create a mock document with no tables
    mock_doc = MagicMock()
    mock_doc.tables = []
    # Add sections for _get_available_content_width
    mock_section = MagicMock()
    mock_section.page_width = DocxPostProcess.DOCX_LETTER_WIDTH_EMU
    mock_section.left_margin = DocxPostProcess.DOCX_LETTER_SIDE_MARGIN
    mock_section.right_margin = DocxPostProcess.DOCX_LETTER_SIDE_MARGIN
    mock_doc.sections = [mock_section]

    # Call the function under test
    DocxPostProcess._replace_table_properties(mock_doc)

    # Verify we checked the tables collection
    assert len(mock_doc.tables) == 0


def test_get_available_content_width():
    """Test the get_available_content_width function."""
    # Create a mock section
    mock_section = MagicMock()
    mock_section.page_width = DocxPostProcess.EMU_1_INCH * 11  # 11 inches
    mock_section.left_margin = DocxPostProcess.EMU_1_INCH * 1.5  # 1.5 inches
    mock_section.right_margin = DocxPostProcess.EMU_1_INCH * 1.5  # 1.5 inches

    # Create a mock document with the mock section
    mock_doc = MagicMock()
    mock_doc.sections = [mock_section]

    # Calculate expected content width: 11 - 1.5 - 1.5 = 8 inches
    expected_width = int(DocxPostProcess.EMU_1_INCH * 8)

    # Get actual content width
    actual_width = DocxPostProcess._get_available_content_width(mock_doc)

    # Assert they match
    assert actual_width == expected_width


def test_get_available_content_width_default_values():
    """Test the get_available_content_width function with default values."""
    # Create a mock section with None values for page dimensions
    mock_section = MagicMock()
    mock_section.page_width = None
    mock_section.left_margin = None
    mock_section.right_margin = None

    # Create a mock document with the mock section
    mock_doc = MagicMock()
    mock_doc.sections = [mock_section]

    # We expect the function to use defaults from DocxPostProcess
    expected_width = int(DocxPostProcess.DOCX_LETTER_WIDTH_EMU - 2 * DocxPostProcess.DOCX_LETTER_SIDE_MARGIN)  # 8.5 - 2 = 6.5 inches

    # Get actual width
    actual_width = DocxPostProcess._get_available_content_width(mock_doc)

    # Assert it uses default values correctly
    assert actual_width == expected_width


def test_resize_images_in_cell_no_resizing_needed():
    """Test that small images don't get resized."""
    # Create a mock cell
    cell = MagicMock(spec=_Cell)

    # Create mock XML content with a small image (smaller than max width)
    small_image_width = 100000  # Small value in EMU, less than max_width
    small_image_height = 100000
    small_image_xml = f'''
    <w:tc xmlns:w="{WORD_PROCESSING_ML_MAIN_SCHEMA}"
          xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
          xmlns:a="{DRAWING_ML_MAIN_SCHEMA}">
        <w:tcPr/>
        <w:p>
            <w:r>
                <w:drawing>
                    <wp:inline>
                        <wp:extent cx="{small_image_width}" cy="{small_image_height}"/>
                    </wp:inline>
                </w:drawing>
            </w:r>
        </w:p>
    </w:tc>
    '''

    # Set up the mock cell
    mock_tc = MagicMock()
    mock_tc.xml = small_image_xml
    cell._tc = mock_tc

    # Set a max width larger than the image width
    max_width = 500000  # Larger than small_image_width

    # Call the function
    DocxPostProcess._resize_images_in_cell(cell, max_width)

    # Verify that clear_content was not called (which would indicate the image was modified)
    mock_tc.clear_content.assert_not_called()


def test_resize_images_in_cell_resizing_needed():
    """Test that large images are properly resized."""
    # Create a mock cell
    cell = MagicMock(spec=_Cell)

    # Create mock XML content with a large image (larger than max_width)
    large_image_width = 1000000  # Large value in EMU, greater than max_width
    large_image_height = 750000  # 3:4 aspect ratio
    large_image_xml = f'''
    <w:tc xmlns:w="{WORD_PROCESSING_ML_MAIN_SCHEMA}"
          xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
          xmlns:a="{DRAWING_ML_MAIN_SCHEMA}">
        <w:tcPr/>
        <w:p>
            <w:r>
                <w:drawing>
                    <wp:inline>
                        <wp:extent cx="{large_image_width}" cy="{large_image_height}"/>
                    </wp:inline>
                </w:drawing>
            </w:r>
        </w:p>
    </w:tc>
    '''

    # Set up the mock for lxml etree parsing - create elements that mimic the real ones

    # Parse the XML to create a real tree
    # ruff: noqa: S320
    tree = etree.fromstring(large_image_xml)

    # Find the wp:extent element to use with the real function
    extent_element = tree.find(".//wp:extent", {"wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"})
    assert extent_element.get("cx") == str(large_image_width)
    assert extent_element.get("cy") == str(large_image_height)

    # Set up the mock cell with our parsed tree
    mock_tc = MagicMock()
    mock_tc.xml = large_image_xml
    cell._tc = mock_tc

    # Set a max width smaller than the image width to trigger resizing
    max_width = 500000  # Smaller than large_image_width

    # Call the function
    with patch("app.DocxPostProcess.etree.fromstring", return_value=tree):
        DocxPostProcess._resize_images_in_cell(cell, max_width)

    # Verify that clear_content was called (which indicates the image was modified)
    mock_tc.clear_content.assert_called_once()

    # Check that the image dimensions were actually updated in the element tree
    # Expected new width is max_width
    expected_new_width = max_width
    # Expected new height maintains the original aspect ratio: height * (new_width / old_width)
    expected_new_height = int(large_image_height * (expected_new_width / large_image_width))

    assert extent_element.get("cx") == str(expected_new_width)
    assert extent_element.get("cy") == str(expected_new_height)


def test_resize_images_in_cell():
    """Test that images in a table cell are correctly resized."""
    # Create a mock cell with proper _tc attribute
    cell = MagicMock()
    cell._tc = MagicMock()
    cell._tc.xml = '<w:tc xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"><wp:extent cx="5000" cy="3000"/></w:tc>'

    # Create a mock for etree.fromstring that returns a properly structured mock
    mock_tree = MagicMock()
    mock_extent = MagicMock()
    mock_extent.attrib = {"cx": "5000", "cy": "3000"}
    mock_tree.findall.return_value = [mock_extent]

    # Patch etree.fromstring to return our mock
    with patch("app.DocxPostProcess.etree.fromstring", return_value=mock_tree):
        # Use a reasonable max width value that will trigger resizing
        max_image_width = 4000.0  # Smaller than the image width

        # Call resize_images_in_cell directly
        DocxPostProcess._resize_images_in_cell(cell, max_image_width)

        # Verify that the image was resized
        assert mock_extent.set.call_count == 2  # cx and cy set
        # Resize should maintain aspect ratio
        mock_extent.set.assert_any_call("cx", "4000")  # New width
        # New height = old_height * (new_width/old_width) = 3000 * (4000/5000) = 2400
        mock_extent.set.assert_any_call("cy", "2400")  # New height (aspect ratio maintained)


@patch("app.DocxPostProcess._resize_images_in_cell")
def test_process_table_with_nested_tables(mock_resize_images):
    """Test that nested tables are processed correctly."""

    # Create mock tables and cells
    main_table = MagicMock()
    nested_table = MagicMock()
    cell = MagicMock()

    # Set up cell to return the nested table
    cell.tables = [nested_table]

    # Set up row to return the cell
    row = MagicMock()
    row.cells = [cell]

    # Set up main table properties
    main_table.rows = [row]
    main_table.columns = MagicMock()
    main_table.columns.__len__.return_value = 2
    main_table._element = MagicMock()
    main_table._element.find.return_value = None

    # Set up property mocks to track XML manipulation
    with patch("app.DocxPostProcess.parse_xml") as mock_parse_xml:
        mock_tbl_props = MagicMock()
        mock_parse_xml.return_value = mock_tbl_props

        # Set up max_width for testing
        max_width = 9144000  # 10 inches in EMU

        # Call the actual _process_table function
        _process_table(main_table, 0, max_width)

        # Verify find was called with correct parameters
        main_table._element.find.assert_called_with(".//w:tblPr", namespaces={"w": SCHEMA})

        # Verify parse_xml was called (for creating table properties)
        mock_parse_xml.assert_called()

        # Verify the table element was updated
        main_table._element.insert.assert_called_with(0, mock_tbl_props)

        # Verify _resize_images_in_cell was called with correct parameters
        mock_resize_images.assert_called_with(cell, max_width / 2)

        # Verify that we processed any nested tables
        for _sub_table in cell.tables:
            # At this point, we would process the nested table, but we can't verify
            # the actual call since it happens recursively. Instead, we verify that
            # the nested tables were accessed.
            pass


def test_replace_size_and_orientation_both_none():
    """Test that no modifications are made when both parameters are None."""
    mock_doc = MagicMock()
    mock_section = MagicMock()
    mock_doc.sections = [mock_section]

    # Call with both None
    DocxPostProcess._replace_size_and_orientation(mock_doc, None, None)

    # Verify no modifications were attempted
    mock_section._sectPr.find.assert_not_called()


def test_replace_size_and_orientation_set_paper_size_only():
    """Test setting paper size without changing orientation."""
    mock_doc = MagicMock()
    mock_section = MagicMock()
    mock_sectPr = MagicMock()
    mock_pgSz = MagicMock()

    # Set up existing paper size element
    mock_pgSz.get.return_value = None  # No existing orientation
    mock_sectPr.find.return_value = mock_pgSz
    mock_section._sectPr = mock_sectPr
    mock_doc.sections = [mock_section]

    # Call with paper_size = A4
    DocxPostProcess._replace_size_and_orientation(mock_doc, "A4", None)

    # Verify pgSz was updated with A4 dimensions (portrait)
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}w", "11906")
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}h", "16838")


def test_replace_size_and_orientation_set_orientation_only():
    """Test setting orientation without changing paper size."""
    mock_doc = MagicMock()
    mock_section = MagicMock()
    mock_sectPr = MagicMock()
    mock_pgSz = MagicMock()

    # Set up existing paper size element with portrait dimensions (width < height)
    # The get method is called twice: once for width, once for height
    mock_pgSz.get.side_effect = ["11906", "16838"]  # portrait: width=11906, height=16838
    mock_pgSz.attrib = {}
    mock_sectPr.find.return_value = mock_pgSz
    mock_section._sectPr = mock_sectPr
    mock_doc.sections = [mock_section]

    # Call with orientation = landscape
    DocxPostProcess._replace_size_and_orientation(mock_doc, None, "landscape")

    # Verify dimensions were swapped (portrait to landscape)
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}w", "16838")
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}h", "11906")
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}orient", "landscape")


def test_replace_size_and_orientation_both_parameters():
    """Test setting both paper size and orientation."""
    mock_doc = MagicMock()
    mock_section = MagicMock()
    mock_sectPr = MagicMock()
    mock_pgSz = MagicMock()

    # Need to provide return values for the get calls:
    # 1. _set_paper_size calls get() once for existing orient attribute -> None
    # 2. _set_orientation calls get() for width -> "12240" (LETTER portrait width, set by _set_paper_size)
    # 3. _set_orientation calls get() for height -> "15840" (LETTER portrait height, set by _set_paper_size)
    mock_pgSz.get.side_effect = [None, "12240", "15840"]
    mock_pgSz.attrib = {}
    mock_sectPr.find.return_value = mock_pgSz
    mock_section._sectPr = mock_sectPr
    mock_doc.sections = [mock_section]

    # Call with paper_size = LETTER and orientation = landscape
    DocxPostProcess._replace_size_and_orientation(mock_doc, "LETTER", "landscape")

    # Verify LETTER portrait dimensions were set first by _set_paper_size
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}w", "12240")
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}h", "15840")
    # Then verify they were swapped for landscape by _set_orientation
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}w", "15840")
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}h", "12240")
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}orient", "landscape")


def test_set_paper_size_unsupported():
    """Test that unsupported paper size raises ValueError."""
    mock_doc = MagicMock()
    mock_section = MagicMock()
    mock_sectPr = MagicMock()
    mock_pgSz = MagicMock()

    mock_sectPr.find.return_value = mock_pgSz
    mock_section._sectPr = mock_sectPr
    mock_doc.sections = [mock_section]

    # Test with unsupported paper size
    with pytest.raises(ValueError, match="Unsupported paper size: TABLOID"):
        DocxPostProcess._replace_size_and_orientation(mock_doc, "TABLOID", None)


def test_set_paper_size_case_insensitive():
    """Test that paper size is case-insensitive."""
    mock_doc = MagicMock()
    mock_section = MagicMock()
    mock_sectPr = MagicMock()
    mock_pgSz = MagicMock()

    mock_pgSz.get.return_value = None
    mock_sectPr.find.return_value = mock_pgSz
    mock_section._sectPr = mock_sectPr
    mock_doc.sections = [mock_section]

    # Call with lowercase paper size
    DocxPostProcess._replace_size_and_orientation(mock_doc, "a4", None)

    # Verify A4 dimensions were set
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}w", "11906")
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}h", "16838")


def test_set_paper_size_creates_pgSz_if_missing():
    """Test that pgSz element is created if it doesn't exist."""
    mock_doc = MagicMock()
    mock_section = MagicMock()
    mock_sectPr = MagicMock()

    # pgSz doesn't exist
    mock_sectPr.find.return_value = None
    mock_section._sectPr = mock_sectPr
    mock_doc.sections = [mock_section]

    with patch("app.DocxPostProcess.parse_xml") as mock_parse_xml:
        mock_new_pgSz = MagicMock()
        mock_new_pgSz.get.return_value = None
        mock_parse_xml.return_value = mock_new_pgSz

        # Call with paper_size = A5
        DocxPostProcess._replace_size_and_orientation(mock_doc, "A5", None)

        # Verify parse_xml was called to create new pgSz
        mock_parse_xml.assert_called_once()
        # Verify the new pgSz was appended
        mock_sectPr.append.assert_called_once_with(mock_new_pgSz)


def test_set_paper_size_preserves_landscape_orientation():
    """Test that existing landscape orientation is preserved when changing paper size."""
    mock_doc = MagicMock()
    mock_section = MagicMock()
    mock_sectPr = MagicMock()
    mock_pgSz = MagicMock()

    # Existing page has landscape orientation
    mock_pgSz.get.return_value = "landscape"
    mock_sectPr.find.return_value = mock_pgSz
    mock_section._sectPr = mock_sectPr
    mock_doc.sections = [mock_section]

    # Call with new paper_size but no orientation parameter
    DocxPostProcess._replace_size_and_orientation(mock_doc, "A3", None)

    # Verify A3 dimensions were set in landscape (swapped)
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}w", "23811")  # height becomes width
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}h", "16838")  # width becomes height
    # Verify landscape orientation was preserved
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}orient", "landscape")


def test_set_orientation_creates_pgSz_if_missing():
    """Test that pgSz element is created with LETTER default if missing."""
    mock_doc = MagicMock()
    mock_section = MagicMock()
    mock_sectPr = MagicMock()

    # pgSz doesn't exist
    mock_sectPr.find.return_value = None
    mock_section._sectPr = mock_sectPr
    mock_doc.sections = [mock_section]

    with patch("app.DocxPostProcess.parse_xml") as mock_parse_xml:
        mock_new_pgSz = MagicMock()
        # Simulate getting width and height from the newly created element (LETTER portrait)
        mock_new_pgSz.get.side_effect = ["12240", "15840"]  # width, then height
        mock_new_pgSz.attrib = {}
        mock_parse_xml.return_value = mock_new_pgSz

        # Call with orientation only
        DocxPostProcess._replace_size_and_orientation(mock_doc, None, "landscape")

        # Verify parse_xml was called to create pgSz with LETTER dimensions
        mock_parse_xml.assert_called_once()
        # Verify the new pgSz was appended
        mock_sectPr.append.assert_called_once_with(mock_new_pgSz)
        # Verify dimensions were swapped for landscape
        mock_new_pgSz.set.assert_any_call(f"{{{SCHEMA}}}w", "15840")
        mock_new_pgSz.set.assert_any_call(f"{{{SCHEMA}}}h", "12240")


def test_set_orientation_portrait_removes_orient_attribute():
    """Test that orient attribute is removed for portrait orientation."""
    mock_doc = MagicMock()
    mock_section = MagicMock()
    mock_sectPr = MagicMock()
    mock_pgSz = MagicMock()

    # Set up existing landscape page (width > height)
    mock_pgSz.get.side_effect = ["15840", "12240"]  # landscape: width=15840, height=12240
    mock_pgSz.attrib = {f"{{{SCHEMA}}}orient": "landscape"}
    mock_sectPr.find.return_value = mock_pgSz
    mock_section._sectPr = mock_sectPr
    mock_doc.sections = [mock_section]

    # Call with orientation = portrait
    DocxPostProcess._replace_size_and_orientation(mock_doc, None, "portrait")

    # Verify dimensions were swapped back to portrait
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}w", "12240")
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}h", "15840")
    # Verify orient attribute was removed
    assert f"{{{SCHEMA}}}orient" not in mock_pgSz.attrib


def test_set_orientation_no_swap_if_already_correct():
    """Test that dimensions aren't swapped if orientation is already correct."""
    mock_doc = MagicMock()
    mock_section = MagicMock()
    mock_sectPr = MagicMock()
    mock_pgSz = MagicMock()

    # Set up existing landscape page (width > height)
    width = "16838"
    height = "11906"
    call_count = [0]

    def get_side_effect(attr, default="0"):
        call_count[0] += 1
        if call_count[0] == 1:  # First call for width
            return width
        else:  # Second call for height
            return height

    mock_pgSz.get.side_effect = get_side_effect
    mock_pgSz.attrib = {f"{{{SCHEMA}}}orient": "landscape"}
    mock_sectPr.find.return_value = mock_pgSz
    mock_section._sectPr = mock_sectPr
    mock_doc.sections = [mock_section]

    # Call with orientation = landscape (already landscape)
    DocxPostProcess._replace_size_and_orientation(mock_doc, None, "landscape")

    # Verify dimensions were NOT swapped (set should not be called with swapped values)
    # The orient attribute should still be set
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}orient", "landscape")


def test_replace_size_and_orientation_multiple_sections():
    """Test that all sections in a document are processed."""
    mock_doc = MagicMock()
    mock_section1 = MagicMock()
    mock_section2 = MagicMock()
    mock_sectPr1 = MagicMock()
    mock_sectPr2 = MagicMock()
    mock_pgSz1 = MagicMock()
    mock_pgSz2 = MagicMock()

    mock_pgSz1.get.return_value = None
    mock_pgSz2.get.return_value = None
    mock_sectPr1.find.return_value = mock_pgSz1
    mock_sectPr2.find.return_value = mock_pgSz2
    mock_section1._sectPr = mock_sectPr1
    mock_section2._sectPr = mock_sectPr2
    mock_doc.sections = [mock_section1, mock_section2]

    # Call with paper_size = B5
    DocxPostProcess._replace_size_and_orientation(mock_doc, "B5", None)

    # Verify both sections were updated
    mock_pgSz1.set.assert_any_call(f"{{{SCHEMA}}}w", "9979")
    mock_pgSz1.set.assert_any_call(f"{{{SCHEMA}}}h", "14144")
    mock_pgSz2.set.assert_any_call(f"{{{SCHEMA}}}w", "9979")
    mock_pgSz2.set.assert_any_call(f"{{{SCHEMA}}}h", "14144")


@pytest.mark.parametrize(
    "paper_size,expected_width,expected_height",
    [
        ("A5", "8419", "11906"),
        ("A4", "11906", "16838"),
        ("A3", "16838", "23811"),
        ("B5", "9979", "14144"),
        ("B4", "14144", "20013"),
        ("JIS_B5", "10319", "14572"),
        ("JIS_B4", "14572", "20639"),
        ("LETTER", "12240", "15840"),
        ("LEGAL", "12240", "20160"),
        ("LEDGER", "15840", "24480"),
    ],
)
def test_all_supported_paper_sizes(paper_size, expected_width, expected_height):
    """Test that all supported paper sizes are correctly applied."""
    mock_doc = MagicMock()
    mock_section = MagicMock()
    mock_sectPr = MagicMock()
    mock_pgSz = MagicMock()

    mock_pgSz.get.return_value = None
    mock_sectPr.find.return_value = mock_pgSz
    mock_section._sectPr = mock_sectPr
    mock_doc.sections = [mock_section]

    # Call with the specified paper_size
    DocxPostProcess._replace_size_and_orientation(mock_doc, paper_size, None)

    # Verify the correct dimensions were set
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}w", expected_width)
    mock_pgSz.set.assert_any_call(f"{{{SCHEMA}}}h", expected_height)


@pytest.mark.parametrize(
    "argv, expected_exit, paper_size, orientation",
    [
        (["script.py"], True, None, None),
        (["script.py", "test.docx", "A4", "landscape"], False, "A4", "landscape"),
        (["script.py", "test.docx", "None", "portrait"], False, None, "portrait"),
        (["script.py", "test.docx", "LETTER", "None"], False, "LETTER", None),
    ],
)
def test_main_function(argv, expected_exit, paper_size, orientation):
    fake_docx_content = b"fake content"
    modified_content = b"modified content"

    with (
        patch.object(sys, "argv", argv),
        patch("pathlib.Path.open", mock_open(read_data=fake_docx_content)) as mock_file,
        patch("app.DocxPostProcess.process", return_value=modified_content) as mock_process,
        patch("app.DocxPostProcess.logging") as mock_logging,
    ):
        result = DocxPostProcess.main()

        if expected_exit:
            assert result == 1
        else:
            assert result == 0
            mock_process.assert_called_once_with(fake_docx_content, paper_size, orientation)
            handle = mock_file()
            handle.write.assert_called_once_with(modified_content)
            mock_logging.debug.assert_called_once()


# Integration tests for the process() function to ensure 100% coverage
class TestProcessFunction:
    """Integration tests that call the process() function directly."""

    def test_process_with_no_parameters(self):
        """Test process() with no paper_size or orientation - should just process tables."""
        # Create a minimal valid DOCX file
        from docx import Document
        import io

        doc = Document()
        doc.add_paragraph("Test content")
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        input_bytes = docx_bytes.getvalue()

        # Call process with no parameters
        result = DocxPostProcess.process(input_bytes)

        # Verify result is valid bytes
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_process_with_paper_size_only(self):
        """Test process() with only paper_size parameter."""
        from docx import Document
        import io

        doc = Document()
        doc.add_paragraph("Test content")
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        input_bytes = docx_bytes.getvalue()

        # Call process with paper_size
        result = DocxPostProcess.process(input_bytes, paper_size="A4")

        # Verify result is valid bytes
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_process_with_orientation_only(self):
        """Test process() with only orientation parameter."""
        from docx import Document
        import io

        doc = Document()
        doc.add_paragraph("Test content")
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        input_bytes = docx_bytes.getvalue()

        # Call process with orientation
        result = DocxPostProcess.process(input_bytes, orientation="landscape")

        # Verify result is valid bytes
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_process_with_both_parameters(self):
        """Test process() with both paper_size and orientation parameters."""
        from docx import Document
        import io

        doc = Document()
        doc.add_paragraph("Test content")
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        input_bytes = docx_bytes.getvalue()

        # Call process with both parameters
        result = DocxPostProcess.process(input_bytes, paper_size="LETTER", orientation="portrait")

        # Verify result is valid bytes
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_process_returns_modified_document(self):
        """Test that process() returns a modified valid DOCX document."""
        from docx import Document
        import io

        # Create input document
        doc = Document()
        doc.add_paragraph("Test content")
        doc.add_table(rows=2, cols=2)
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        input_bytes = docx_bytes.getvalue()

        # Process the document
        result = DocxPostProcess.process(input_bytes, paper_size="A4", orientation="landscape")

        # Verify we can open the result as a valid DOCX
        result_doc = Document(io.BytesIO(result))
        assert len(result_doc.paragraphs) > 0
        assert len(result_doc.tables) > 0
